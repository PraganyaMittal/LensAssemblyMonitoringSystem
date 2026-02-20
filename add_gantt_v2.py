"""
1) Opens SRS_FactoryMonitoring_Final 3 (1).docx
   - Inserts Gantt chart tables after SRS and NFR tables
   - No "X" markers, no color in week cells - just light grey fill for scheduled weeks
   - Table sized to fit one horizontal line in landscape
   - Saves as SRS_FactoryMonitoring_Final_WithGantt_v2.docx

2) Creates SRS_Gantt_Charts.xlsx with both Gantt charts as Excel sheets
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml, OxmlElement
import os

# openpyxl for Excel
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

BASE = os.path.dirname(os.path.abspath(__file__))
SRC = os.path.join(BASE, "SRS_FactoryMonitoring_Final 3 (1).docx")
OUT_DOCX = os.path.join(BASE, "SRS_FactoryMonitoring_Final_WithGantt_v4.docx")
OUT_XLSX = os.path.join(BASE, "SRS_Gantt_Charts_v5.xlsx")

MONTHS = ["Mar'26", "Apr'26", "May'26", "Jun'26"]
WEEKS_PER_MONTH = 4
TOTAL_WEEKS = 16

# Gantt schedule data
FR_TASKS = [
    ("SRS-001", "Real-Time Dashboard & Monitoring", "Dashboard UI, status indicators, filtering", 0, 3),
    ("SRS-002", "PC Details & Administration", "PC detail views, metadata editing", 2, 5),
    ("SRS-003", "PC Model & Config Operations", "Upload/download/change models & configs", 3, 6),
    ("SRS-004", "Model Library & Distribution", "Library page, distribution pipeline", 4, 8),
    ("SRS-005", "Model Editor & XML Visual Editor", "Code editor, XML param editor, diff view", 5, 9),
    ("SRS-006", "Version History & Rollback", "Auto-versioning, diff viewer, revert", 7, 10),
    ("SRS-007", "Log Analyser & Visualization", "Log parsing, Gantt charts, images", 3, 8),
    ("SRS-008", "Agent Registration & Heartbeat", "Registration dialog, heartbeat loop", 0, 3),
    ("SRS-009", "Agent Sync Operations", "Model/config/log sync, image upload", 2, 6),
    ("SRS-010", "UI Framework & Navigation", "Sidebar, themes, routing, 404 page", 0, 4),
]

NFR_TASKS = [
    ("NFR-001", "Performance & Reliability", "Load time, caching, timeouts, offline detection", 0, 15),
    ("NFR-002", "Usability & Security", "Themes, color-coding, path validation", 0, 15),
]

# ============================================================
#  PART 1: WORD DOCUMENT
# ============================================================

def shade(cell, color):
    cell._tc.get_or_add_tcPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>'))

def set_cell_text(cell, text, size=7, bold=False, align=None, font_color=None):
    cell.text = text
    for p in cell.paragraphs:
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        if align:
            p.alignment = align
        for r in p.runs:
            r.font.size = Pt(size)
            r.font.bold = bold
            r.font.name = 'Calibri'
            if font_color:
                rPr = r._element.get_or_add_rPr()
                c_elem = OxmlElement('w:color')
                c_elem.set(qn('w:val'), font_color)
                rPr.append(c_elem)

def build_gantt_table_docx(doc, title_text, tasks):
    """Build Gantt table in doc. Returns (title_paragraph_xml, table_element)."""
    # Title paragraph
    title_para = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '240')
    spacing.set(qn('w:after'), '120')
    pPr.append(spacing)
    title_para.append(pPr)
    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rPr.append(OxmlElement('w:b'))
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '22')
    rPr.append(sz)
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '1F4E79')
    rPr.append(color)
    run.append(rPr)
    t_elem = OxmlElement('w:t')
    t_elem.text = title_text
    run.append(t_elem)
    title_para.append(run)

    num_cols = 3 + TOTAL_WEEKS
    num_rows = 2 + len(tasks)

    tbl = doc.add_table(rows=num_rows, cols=num_cols)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False

    # Row 0: Main headers
    row0 = tbl.rows[0]
    for ci, hdr in enumerate(["SRS ID", "Feature / Module", "Sub Tasks"]):
        set_cell_text(row0.cells[ci], hdr, size=8, bold=True,
                      align=WD_ALIGN_PARAGRAPH.CENTER, font_color='FFFFFF')
        shade(row0.cells[ci], "1F4E79")

    for mi, month in enumerate(MONTHS):
        sc = 3 + mi * WEEKS_PER_MONTH
        ec = sc + WEEKS_PER_MONTH - 1
        merged = row0.cells[sc]
        merged.merge(row0.cells[ec])
        set_cell_text(merged, month, size=8, bold=True,
                      align=WD_ALIGN_PARAGRAPH.CENTER, font_color='FFFFFF')
        shade(merged, "1F4E79")

    # Row 1: Week sub-headers
    row1 = tbl.rows[1]
    for ci in range(3):
        set_cell_text(row1.cells[ci], "", size=7)
        shade(row1.cells[ci], "D6E4F0")
    for wi in range(TOTAL_WEEKS):
        cell = row1.cells[3 + wi]
        set_cell_text(cell, f"W{(wi % 4) + 1}", size=7, bold=True,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
        shade(cell, "D6E4F0")

    # Data rows
    for ti, (srs_id, feature, sub_task, sw, ew) in enumerate(tasks):
        row = tbl.rows[2 + ti]
        bg = "F2F7FB" if ti % 2 == 0 else "FFFFFF"
        set_cell_text(row.cells[0], srs_id, size=8, bold=True)
        set_cell_text(row.cells[1], feature, size=7)
        set_cell_text(row.cells[2], sub_task, size=7)
        for ci in range(3):
            shade(row.cells[ci], bg)

        for wi in range(TOTAL_WEEKS):
            cell = row.cells[3 + wi]
            set_cell_text(cell, "", size=6)
            # All week cells plain white - no fill at all

    # Column widths - landscape page ~27.9cm usable (11in - 2*1.5cm margins)
    # 3 text cols + 16 week cols
    for row in tbl.rows:
        row.cells[0].width = Cm(2.0)
        row.cells[1].width = Cm(4.5)
        row.cells[2].width = Cm(5.5)
        for wi in range(TOTAL_WEEKS):
            row.cells[3 + wi].width = Cm(0.95)

    tbl_element = tbl._element
    return title_para, tbl_element


def generate_word():
    doc = Document(SRC)
    body = doc.element.body

    srs_tbl_elem = doc.tables[7]._element  # Functional SRS table
    nfr_tbl_elem = doc.tables[8]._element  # NFR table

    title1, gantt1 = build_gantt_table_docx(doc,
        "Functional SRS — Work Distribution Schedule (March – June 2026)", FR_TASKS)
    title2, gantt2 = build_gantt_table_docx(doc,
        "Non-Functional SRS — Work Distribution Schedule (March – June 2026)", NFR_TASKS)

    body.remove(gantt1)
    body.remove(gantt2)

    # Spacer paragraphs
    def make_spacer():
        sp = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:after'), '60')
        pPr.append(spacing)
        sp.append(pPr)
        return sp

    # Insert after SRS table: title -> gantt -> spacer
    srs_tbl_elem.addnext(make_spacer())
    srs_tbl_elem.addnext(gantt1)
    srs_tbl_elem.addnext(title1)

    # Insert after NFR table: title -> gantt -> spacer
    nfr_tbl_elem.addnext(make_spacer())
    nfr_tbl_elem.addnext(gantt2)
    nfr_tbl_elem.addnext(title2)

    doc.save(OUT_DOCX)
    print(f"Word → {OUT_DOCX}")


# ============================================================
#  PART 2: EXCEL FILE
# ============================================================

def create_gantt_sheet(wb, sheet_name, title, tasks):
    ws = wb.create_sheet(title=sheet_name)

    # Styles
    hdr_fill = PatternFill(start_color="1F4E79", end_color="1F4E79", fill_type="solid")
    hdr_font = Font(name="Calibri", size=9, bold=True, color="FFFFFF")
    week_hdr_fill = PatternFill(start_color="D6E4F0", end_color="D6E4F0", fill_type="solid")
    week_hdr_font = Font(name="Calibri", size=8, bold=True)
    data_font = Font(name="Calibri", size=9)
    data_bold_font = Font(name="Calibri", size=9, bold=True)
    bar_fill = PatternFill(start_color="BDD7EE", end_color="BDD7EE", fill_type="solid")
    white_fill = PatternFill(start_color="FFFFFF", end_color="FFFFFF", fill_type="solid")
    center = Alignment(horizontal="center", vertical="center", wrap_text=True)
    left_wrap = Alignment(horizontal="left", vertical="center", wrap_text=True)
    thin_border = Border(
        left=Side(style="thin"), right=Side(style="thin"),
        top=Side(style="thin"), bottom=Side(style="thin"),
    )

    # Title row
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=3 + TOTAL_WEEKS)
    title_cell = ws.cell(row=1, column=1, value=title)
    title_cell.font = Font(name="Calibri", size=12, bold=True, color="1F4E79")
    title_cell.alignment = Alignment(horizontal="left", vertical="center")
    ws.row_dimensions[1].height = 25

    # Row 2: Main headers
    row_num = 2
    headers = ["SRS ID", "Feature / Module", "Sub Tasks"]
    for ci, h in enumerate(headers, 1):
        cell = ws.cell(row=row_num, column=ci, value=h)
        cell.fill = hdr_fill
        cell.font = hdr_font
        cell.alignment = center
        cell.border = thin_border

    for mi, month in enumerate(MONTHS):
        sc = 4 + mi * WEEKS_PER_MONTH
        ec = sc + WEEKS_PER_MONTH - 1
        ws.merge_cells(start_row=row_num, start_column=sc, end_row=row_num, end_column=ec)
        cell = ws.cell(row=row_num, column=sc, value=month)
        cell.fill = hdr_fill
        cell.font = hdr_font
        cell.alignment = center
        cell.border = thin_border
        # Border all merged cells
        for c in range(sc, ec + 1):
            ws.cell(row=row_num, column=c).border = thin_border

    ws.row_dimensions[row_num].height = 22

    # Row 3: Week sub-headers
    row_num = 3
    for ci in range(1, 4):
        cell = ws.cell(row=row_num, column=ci, value="")
        cell.fill = week_hdr_fill
        cell.border = thin_border
    for wi in range(TOTAL_WEEKS):
        cell = ws.cell(row=row_num, column=4 + wi, value=f"W{(wi % 4) + 1}")
        cell.fill = week_hdr_fill
        cell.font = week_hdr_font
        cell.alignment = center
        cell.border = thin_border

    ws.row_dimensions[row_num].height = 18

    # Data rows
    for ti, (srs_id, feature, sub_task, sw, ew) in enumerate(tasks):
        row_num = 4 + ti

        c1 = ws.cell(row=row_num, column=1, value=srs_id)
        c1.font = data_bold_font
        c1.alignment = left_wrap
        c1.fill = white_fill
        c1.border = thin_border

        c2 = ws.cell(row=row_num, column=2, value=feature)
        c2.font = data_font
        c2.alignment = left_wrap
        c2.fill = white_fill
        c2.border = thin_border

        c3 = ws.cell(row=row_num, column=3, value=sub_task)
        c3.font = data_font
        c3.alignment = left_wrap
        c3.fill = white_fill
        c3.border = thin_border

        for wi in range(TOTAL_WEEKS):
            cell = ws.cell(row=row_num, column=4 + wi, value="")
            cell.border = thin_border
            cell.alignment = center
            cell.fill = white_fill  # every week cell plain white, no color

        ws.row_dimensions[row_num].height = 28

    # Column widths
    ws.column_dimensions[get_column_letter(1)].width = 12
    ws.column_dimensions[get_column_letter(2)].width = 32
    ws.column_dimensions[get_column_letter(3)].width = 38
    for wi in range(TOTAL_WEEKS):
        ws.column_dimensions[get_column_letter(4 + wi)].width = 5

    # Freeze panes
    ws.freeze_panes = "D4"


def generate_excel():
    wb = Workbook()
    # Remove default sheet
    wb.remove(wb.active)

    create_gantt_sheet(wb, "Functional SRS",
        "Functional SRS — Work Distribution Schedule (March – June 2026)", FR_TASKS)
    create_gantt_sheet(wb, "Non-Functional SRS",
        "Non-Functional SRS — Work Distribution Schedule (March – June 2026)", NFR_TASKS)

    wb.save(OUT_XLSX)
    print(f"Excel → {OUT_XLSX}")


if __name__ == "__main__":
    generate_word()
    generate_excel()
    print("All done!")
