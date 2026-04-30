import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as patches
import numpy as np
import os

OUT = os.path.dirname(os.path.abspath(__file__))

def arrow(ax, x1, y1, x2, y2, txt, fs=9, side='right'):
    ax.annotate('', xy=(x1,y1), xytext=(x2,y2),
        arrowprops=dict(arrowstyle='->', color='#e63946', lw=1.5))
    ha = 'left' if side=='right' else 'right'
    ax.text(x2, y2, f'  {txt}  ', fontsize=fs, fontweight='bold',
        color='#1d3557', ha=ha, va='center',
        bbox=dict(boxstyle='round,pad=0.3', facecolor='#f1faee', edgecolor='#457b9d', lw=1))

# 1. LENS TRAY
def draw_lens_tray():
    fig, ax = plt.subplots(figsize=(10, 7))
    ax.set_xlim(-1, 12); ax.set_ylim(-1, 9); ax.set_aspect('equal')
    ax.add_patch(patches.FancyBboxPatch((-0.3, -0.3), 11.6, 8.6, boxstyle='round,pad=0.2',
        facecolor='#e8e4df', edgecolor='#333', lw=2))
    for i in range(8):
        for j in range(6):
            ax.add_patch(plt.Circle((1+i*1.4, 1+j*1.2), 0.4, fc='#caf0f8', ec='#457b9d', lw=1.5))
    # Param 1: Hole Count
    ax.annotate('', xy=(1,7.8), xytext=(10.8,7.8), arrowprops=dict(arrowstyle='<->', color='#e63946', lw=2))
    ax.text(5.9, 8.2, 'Tray Hole Cnt X = 8', fontsize=11, fontweight='bold', ha='center',
        color='#1d3557', bbox=dict(boxstyle='round,pad=0.3', fc='#f1faee', ec='#457b9d'))
    ax.annotate('', xy=(-0.5,1), xytext=(-0.5,7), arrowprops=dict(arrowstyle='<->', color='#e63946', lw=2))
    ax.text(-0.5, 4, 'Hole Cnt Y = 6', fontsize=10, fontweight='bold', ha='center', rotation=90,
        color='#1d3557', bbox=dict(boxstyle='round,pad=0.3', fc='#f1faee', ec='#457b9d'))
    # Param 2: Interval
    ax.annotate('', xy=(1,0.3), xytext=(2.4,0.3), arrowprops=dict(arrowstyle='<->', color='#2a9d8f', lw=2))
    ax.text(1.7, -0.1, 'Interval X (mm)', fontsize=9, fontweight='bold', ha='center', color='#264653',
        bbox=dict(boxstyle='round,pad=0.2', fc='#e9f5db', ec='#2a9d8f'))
    # Param 3: 1st Hole Pos
    ax.plot(1, 1, 'x', color='#e63946', markersize=15, mew=3)
    ax.annotate('1st Hole Start\nPos (X, Y)', xy=(1,1), xytext=(3.5, -0.5),
        fontsize=10, fontweight='bold', color='#e63946',
        arrowprops=dict(arrowstyle='->', color='#e63946', lw=1.5),
        bbox=dict(boxstyle='round,pad=0.3', fc='#ffe8e8', ec='#e63946'))
    ax.set_title('Lens Tray — Top View', fontsize=16, fontweight='bold', color='#1d3557', pad=15)
    ax.axis('off')
    fig.tight_layout()
    fig.savefig(os.path.join(OUT, 'img_lens_tray.png'), dpi=150, bbox_inches='tight')
    plt.close()

# 2. BARREL TRAY
def draw_barrel_tray():
    fig, ax = plt.subplots(figsize=(10, 7))
    ax.set_xlim(-1, 10); ax.set_ylim(-1, 8); ax.set_aspect('equal')
    ax.add_patch(patches.FancyBboxPatch((-0.3, -0.3), 9.6, 7.6, boxstyle='round,pad=0.2',
        facecolor='#f5ebe0', edgecolor='#333', lw=2))
    for i in range(4):
        for j in range(3):
            x, y = 1+i*2.2, 1.5+j*2
            ax.add_patch(plt.Circle((x,y), 0.8, fc='#d5c4a1', ec='#6b4226', lw=2))
            ax.add_patch(plt.Circle((x,y), 0.6, fc='#fff8e7', ec='#6b4226', lw=1))
    ax.annotate('', xy=(1,6.8), xytext=(7.6,6.8), arrowprops=dict(arrowstyle='<->', color='#e63946', lw=2))
    ax.text(4.3, 7.2, 'Tray Hole Cnt X = 4', fontsize=11, fontweight='bold', ha='center',
        color='#1d3557', bbox=dict(boxstyle='round,pad=0.3', fc='#f1faee', ec='#457b9d'))
    ax.annotate('', xy=(-0.5,1.5), xytext=(-0.5,5.5), arrowprops=dict(arrowstyle='<->', color='#e63946', lw=2))
    ax.text(-0.5, 3.5, 'Hole Cnt Y = 3', fontsize=10, fontweight='bold', ha='center', rotation=90,
        color='#1d3557', bbox=dict(boxstyle='round,pad=0.3', fc='#f1faee', ec='#457b9d'))
    ax.annotate('', xy=(1,0.5), xytext=(3.2,0.5), arrowprops=dict(arrowstyle='<->', color='#2a9d8f', lw=2))
    ax.text(2.1, 0.1, 'Interval X (mm)', fontsize=9, fontweight='bold', ha='center', color='#264653',
        bbox=dict(boxstyle='round,pad=0.2', fc='#e9f5db', ec='#2a9d8f'))
    ax.plot(1, 1.5, 'x', color='#e63946', markersize=15, mew=3)
    ax.annotate('1st Hole Start\nPos (X, Y)', xy=(1,1.5), xytext=(4, 0.1),
        fontsize=10, fontweight='bold', color='#e63946',
        arrowprops=dict(arrowstyle='->', color='#e63946', lw=1.5),
        bbox=dict(boxstyle='round,pad=0.3', fc='#ffe8e8', ec='#e63946'))
    # barcode
    ax.add_patch(patches.Rectangle((7, -0.1), 1.5, 0.5, fc='white', ec='#333', lw=1))
    for k in range(8):
        ax.add_patch(patches.Rectangle((7.1+k*0.15, 0), 0.08, 0.3, fc='black'))
    ax.text(7.75, -0.25, 'Barcode', fontsize=7, ha='center', color='#666')
    ax.set_title('Barrel Tray (Assy Tray) — Top View', fontsize=16, fontweight='bold', color='#1d3557', pad=15)
    ax.axis('off')
    fig.tight_layout()
    fig.savefig(os.path.join(OUT, 'img_barrel_tray.png'), dpi=150, bbox_inches='tight')
    plt.close()

# 3. BARREL CROSS-SECTION
def draw_barrel():
    fig, ax = plt.subplots(figsize=(8, 12))
    ax.set_xlim(-4, 8); ax.set_ylim(-1, 18); ax.set_aspect('equal')
    # Barrel wall (left half cut-away)
    bw = 0.8
    wall_x = [-2.5-bw, -2.5-bw, 2.5+bw, 2.5+bw, 2.5, 2.5, -2.5, -2.5]
    wall_y = [0, 16, 16, 0, 0, 15.5, 15.5, 0]
    ax.fill(wall_x, wall_y, fc='#8d99ae', ec='#2b2d42', lw=2)
    ax.fill([-2.5,-2.5,2.5,2.5], [0,15.5,15.5,0], fc='#edf2f4', ec='none')
    # Lenses and spacers
    elements = [
        ('SP0', 'spacer', 0.5), ('L1', 'lens', 1.5), ('L2', 'lens', 3.0),
        ('SP2', 'spacer', 4.2), ('L3', 'lens', 5.0), ('SP3', 'spacer', 6.5),
        ('L4', 'lens', 7.3), ('SP4', 'spacer', 8.8), ('L5', 'lens', 9.5),
        ('SP5', 'spacer', 11.0), ('L6', 'lens', 11.7), ('SP6', 'spacer', 13.2),
        ('L7', 'lens', 13.8), ('Ring', 'cover', 15.2)
    ]
    for name, typ, y in elements:
        if typ == 'lens':
            t = np.linspace(-2.2, 2.2, 100)
            top = y + 0.8 + 0.3*np.cos(t*np.pi/4.4)
            bot = y + 0.2 - 0.15*np.cos(t*np.pi/4.4)
            ax.fill_between(t, bot, top, fc='#a8dadc', ec='#457b9d', lw=1.5, alpha=0.85)
            ax.text(2.8, y+0.5, name, fontsize=9, fontweight='bold', color='#1d3557', va='center')
        elif typ == 'spacer':
            ax.add_patch(patches.Rectangle((-2.3, y), 4.6, 0.25, fc='#ddd', ec='#999', lw=1))
            ax.text(2.8, y+0.12, name, fontsize=8, color='#666', va='center')
        else:
            ax.add_patch(patches.Rectangle((-2.5, y), 5, 0.4, fc='#b5838d', ec='#6d4c5e', lw=1.5))
            ax.text(2.8, y+0.2, name, fontsize=9, fontweight='bold', color='#6d4c5e', va='center')
    # Param 1: Inner Diameter
    ax.annotate('', xy=(-2.5, -0.3), xytext=(2.5, -0.3), arrowprops=dict(arrowstyle='<->', color='#e63946', lw=2))
    ax.text(0, -0.8, 'Inner Diameter\n(Ø11.300 mm)', fontsize=10, fontweight='bold', ha='center',
        color='#1d3557', bbox=dict(boxstyle='round,pad=0.3', fc='#f1faee', ec='#457b9d'))
    # Param 2: TTL
    ax.annotate('', xy=(-3.8, 0.5), xytext=(-3.8, 15.2), arrowprops=dict(arrowstyle='<->', color='#e63946', lw=2))
    ax.text(-3.8, 8, 'TTL\n8.430mm', fontsize=10, fontweight='bold', ha='center', rotation=90,
        color='#1d3557', bbox=dict(boxstyle='round,pad=0.3', fc='#f1faee', ec='#457b9d'))
    # Param 3: Outer Diameter
    ax.annotate('', xy=(-3.3, 16.5), xytext=(3.3, 16.5), arrowprops=dict(arrowstyle='<->', color='#2a9d8f', lw=2))
    ax.text(0, 17.2, 'Outer Diameter\n(Ø12.300 mm)', fontsize=10, fontweight='bold', ha='center',
        color='#264653', bbox=dict(boxstyle='round,pad=0.3', fc='#e9f5db', ec='#2a9d8f'))
    ax.set_title('Barrel Assembly — Cross Section', fontsize=16, fontweight='bold', color='#1d3557', pad=15)
    ax.axis('off')
    fig.tight_layout()
    fig.savefig(os.path.join(OUT, 'img_barrel.png'), dpi=150, bbox_inches='tight')
    plt.close()

# 4. LENS DETAIL
def draw_lens():
    fig, ax = plt.subplots(figsize=(10, 7))
    ax.set_xlim(-5, 7); ax.set_ylim(-3, 4); ax.set_aspect('equal')
    t = np.linspace(-2.5, 2.5, 200)
    s1 = 0.85 * np.cos(t * np.pi / 5.0)
    s2 = -0.36 * np.cos(t * np.pi / 5.0)
    ax.fill_between(t, s2, s1, fc='#a8dadc', ec='#457b9d', lw=2, alpha=0.8)
    # Param 1: S1 SAG
    ax.annotate('', xy=(0, 0.85), xytext=(0, 0), arrowprops=dict(arrowstyle='<->', color='#e63946', lw=2))
    ax.annotate('S1 SAG\n(0.8500 mm)', xy=(0, 0.42), xytext=(3.5, 2.5),
        fontsize=11, fontweight='bold', color='#1d3557',
        arrowprops=dict(arrowstyle='->', color='#e63946', lw=1.5),
        bbox=dict(boxstyle='round,pad=0.4', fc='#f1faee', ec='#457b9d'))
    # Param 2: Center Thickness
    ax.plot([-0.15,-0.15], [-0.36, 0.85], '--', color='#2a9d8f', lw=1)
    ax.annotate('Center Thickness\n(1.0850 mm)', xy=(-0.15, 0.25), xytext=(-4.5, 2.5),
        fontsize=11, fontweight='bold', color='#264653',
        arrowprops=dict(arrowstyle='->', color='#2a9d8f', lw=1.5),
        bbox=dict(boxstyle='round,pad=0.4', fc='#e9f5db', ec='#2a9d8f'))
    # Param 3: Outer Diameter
    ax.annotate('', xy=(-2.5, -2), xytext=(2.5, -2), arrowprops=dict(arrowstyle='<->', color='#e63946', lw=2))
    ax.text(0, -2.6, 'Outer Diameter (Ø6.020 mm)', fontsize=11, fontweight='bold', ha='center',
        color='#1d3557', bbox=dict(boxstyle='round,pad=0.3', fc='#f1faee', ec='#457b9d'))
    ax.text(2.7, 0.6, 'S1', fontsize=10, color='#457b9d', fontweight='bold')
    ax.text(2.7, -0.3, 'S2', fontsize=10, color='#457b9d', fontweight='bold')
    ax.set_title('Lens — Cross Section (e.g. LENS1: APEL5514ML)', fontsize=14, fontweight='bold', color='#1d3557', pad=15)
    ax.axis('off')
    fig.tight_layout()
    fig.savefig(os.path.join(OUT, 'img_lens.png'), dpi=150, bbox_inches='tight')
    plt.close()

# 5. SPACER DETAIL
def draw_spacer():
    fig, ax = plt.subplots(figsize=(10, 5))
    ax.set_xlim(-5, 7); ax.set_ylim(-2, 3); ax.set_aspect('equal')
    # Spacer ring cross-section (two rectangles = ring profile)
    ax.add_patch(patches.Rectangle((-2.5, 0), 1.3, 0.3, fc='#ddd', ec='#666', lw=2))
    ax.add_patch(patches.Rectangle((1.2, 0), 1.3, 0.3, fc='#ddd', ec='#666', lw=2))
    ax.plot([-1.2, -1.2, 1.2, 1.2], [0, 0.3, 0.3, 0], '--', color='#999', lw=1)
    ax.text(0, 0.15, '(open center)', fontsize=8, ha='center', va='center', color='#999', style='italic')
    # Param 1: Outer Diameter
    ax.annotate('', xy=(-2.5, -0.8), xytext=(2.5, -0.8), arrowprops=dict(arrowstyle='<->', color='#e63946', lw=2))
    ax.text(0, -1.3, 'Outer Diameter (Ø5.800 mm)', fontsize=11, fontweight='bold', ha='center',
        color='#1d3557', bbox=dict(boxstyle='round,pad=0.3', fc='#f1faee', ec='#457b9d'))
    # Param 2: Inner Diameter
    ax.annotate('', xy=(-1.2, 0.7), xytext=(1.2, 0.7), arrowprops=dict(arrowstyle='<->', color='#2a9d8f', lw=2))
    ax.text(0, 1.2, 'Inner Diameter (Ø4.400 mm)', fontsize=11, fontweight='bold', ha='center',
        color='#264653', bbox=dict(boxstyle='round,pad=0.3', fc='#e9f5db', ec='#2a9d8f'))
    # Param 3: Thickness
    ax.annotate('Thickness\n(0.235 mm)', xy=(2.5, 0.15), xytext=(4.5, 1.5),
        fontsize=11, fontweight='bold', color='#1d3557',
        arrowprops=dict(arrowstyle='->', color='#e63946', lw=1.5),
        bbox=dict(boxstyle='round,pad=0.4', fc='#f1faee', ec='#457b9d'))
    ax.annotate('', xy=(2.6, 0), xytext=(2.6, 0.3), arrowprops=dict(arrowstyle='<->', color='#e63946', lw=2))
    ax.set_title('Spacer Ring — Cross Section (FILM type)', fontsize=14, fontweight='bold', color='#1d3557', pad=15)
    ax.axis('off')
    fig.tight_layout()
    fig.savefig(os.path.join(OUT, 'img_spacer.png'), dpi=150, bbox_inches='tight')
    plt.close()

# GENERATE ALL
print("Generating diagrams...")
draw_lens_tray(); print("  [1/5] Lens Tray done")
draw_barrel_tray(); print("  [2/5] Barrel Tray done")
draw_barrel(); print("  [3/5] Barrel done")
draw_lens(); print("  [4/5] Lens done")
draw_spacer(); print("  [5/5] Spacer done")

# BUILD WORD DOCUMENT
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# Title Page
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for _ in range(6): p.add_run('\n')
r = p.add_run('Model Editor\n')
r.font.size = Pt(36); r.bold = True; r.font.color.rgb = RGBColor(0x1d, 0x35, 0x57)
r = p.add_run('UI Design Document\n')
r.font.size = Pt(24); r.font.color.rgb = RGBColor(0x45, 0x7b, 0x9d)
r = p.add_run('\nCAD Base Parameters → SpecData Visual Editor\n')
r.font.size = Pt(14); r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
r = p.add_run('\n\nPrepared for: Client Review\n')
r.font.size = Pt(12); r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

pages = [
    ('img_lens_tray.png', 'Lens Tray — Top View',
     'Each machine has its own lens tray containing one type of lens.\n'
     'The tray is a grid of holes. Key base parameters:\n'
     '• Tray Hole Count (X × Y) — number of holes in the grid\n'
     '• Tray Hole Interval (X, Y) — spacing between holes in mm\n'
     '• 1st Hole Start Position (X, Y) — coordinates of the first hole'),
    ('img_barrel_tray.png', 'Barrel Tray (Assy Tray) — Top View',
     'The barrel tray holds multiple barrels in a grid and moves on the conveyor belt.\n'
     'Each tray has a barcode for identification. Key base parameters:\n'
     '• Tray Hole Count (X × Y) — number of barrel positions\n'
     '• Tray Hole Interval (X, Y) — spacing between barrels in mm\n'
     '• 1st Hole Start Position (X, Y) — coordinates of the first barrel'),
    ('img_barrel.png', 'Barrel Assembly — Cross Section',
     'The barrel holds stacked lenses and spacers. This cross-section shows all\n'
     'components from bottom to top. Each machine inserts one lens + one spacer.\n'
     'Key base parameters from CAD:\n'
     '• Inner Diameter — internal bore diameter\n'
     '• Outer Diameter — external barrel diameter\n'
     '• TTL (Total Track Length) — total depth of the optical stack'),
    ('img_lens.png', 'Lens — Cross Section',
     'Individual lens detail showing optical surfaces S1 (top) and S2 (bottom).\n'
     'Each lens position uses a different material and has unique dimensions.\n'
     'Key base parameters from CAD:\n'
     '• S1 SAG — depth of top surface curvature\n'
     '• Center Thickness — thickness at optical center\n'
     '• Outer Diameter — full lens diameter'),
    ('img_spacer.png', 'Spacer Ring — Cross Section',
     'Spacer rings separate lenses inside the barrel. Each position may have\n'
     'different spacer dimensions. Some positions use multiple sub-spacers.\n'
     'Key base parameters from CAD:\n'
     '• Outer Diameter — spacer outer diameter\n'
     '• Inner Diameter — spacer inner diameter (clear aperture)\n'
     '• Thickness — spacer height')
]

for img_file, title, desc in pages:
    doc.add_page_break()
    h = doc.add_heading(title, level=1)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1d, 0x35, 0x57)
    img_path = os.path.join(OUT, img_file)
    w = Inches(6) if 'barrel.' not in img_file else Inches(4.5)
    doc.add_picture(img_path, width=w)
    last_p = doc.paragraphs[-1]
    last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('')
    p = doc.add_paragraph(desc)
    p.style.font.size = Pt(11)

out_path = os.path.join(OUT, 'Model_Editor_UI_Design.docx')
doc.save(out_path)
print(f"\nWord document saved: {out_path}")
