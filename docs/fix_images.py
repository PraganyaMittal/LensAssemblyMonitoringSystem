import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as patches
import numpy as np
import os

OUT = os.path.dirname(os.path.abspath(__file__))

# BARREL - Stepped interior (staircase), cylindrical outer, ONE lens highlighted
# FLIPPED: Wide=top (open side), Narrow=bottom (closed side)
# Insertion from top, first element goes deepest (bottom), cap at top
def draw_barrel():
    fig, ax = plt.subplots(figsize=(10, 16))
    ax.set_xlim(-6.5, 11); ax.set_ylim(-3, 19.5); ax.set_aspect('equal')

    outer_w = 3.5  # outer wall half-width at widest (top)
    barrel_bot = 0.0
    barrel_top = 14.8

    # FLIPPED: Bottom = narrowest (closed), Top = widest (open)
    # (name, type, y_pos, step_height, inner_half_width)
    # inner_half_width INCREASES going up (wider at top)
    steps = [
        ('SP0',  'spacer', 0.4,  0.6,  1.70),   # Bottom — deepest, narrowest, first inserted
        ('L1',   'lens',   1.0,  1.4,  1.75),
        ('L2',   'lens',   2.4,  1.4,  1.85),
        ('SP2',  'spacer', 3.8,  0.6,  1.95),
        ('L3',   'lens',   4.4,  1.4,  2.05),   # HIGHLIGHTED
        ('SP3',  'spacer', 5.8,  0.6,  2.15),
        ('L4',   'lens',   6.4,  1.4,  2.25),
        ('SP4',  'spacer', 7.8,  0.6,  2.35),
        ('L5',   'lens',   8.4,  1.4,  2.45),
        ('SP5',  'spacer', 9.8,  0.6,  2.55),
        ('L6',   'lens',  10.4,  1.4,  2.65),
        ('SP6',  'spacer', 11.8, 0.6,  2.75),
        ('L7',   'lens',  12.4,  1.4,  2.85),
        ('Ring', 'cover', 13.8,  0.5,  2.95),    # Top — widest, last inserted (cap)
    ]

    # Build staircase inner wall profile
    left_x = []
    left_y = []
    right_x = []
    right_y = []

    left_x.append(-steps[0][4])
    left_y.append(barrel_bot)

    for i, (name, typ, y, h, w) in enumerate(steps):
        left_x.append(-w); left_y.append(y)
        left_x.append(-w); left_y.append(y + h)
        right_x.append(w); right_y.append(y)
        right_x.append(w); right_y.append(y + h)
        if i < len(steps) - 1:
            next_w = steps[i+1][4]
            left_x.append(-next_w); left_y.append(y + h)
            right_x.append(next_w); right_y.append(y + h)

    left_x.append(-steps[-1][4]); left_y.append(barrel_top)
    right_x.append(steps[-1][4]); right_y.append(barrel_top)

    # Outer walls — tapered slightly (wider at top to match inner widening)
    outer_bot_w = 2.2   # outer wall half-width at bottom
    outer_top_w = 3.5   # outer wall half-width at top

    # Left outer wall (trapezoid)
    left_wall = plt.Polygon([
        (-outer_bot_w, barrel_bot), (-outer_top_w, barrel_top),
        (-steps[-1][4], barrel_top), (-steps[0][4], barrel_bot)
    ], fc='#8d99ae', ec='#2b2d42', lw=2)
    ax.add_patch(left_wall)
    # Right outer wall (trapezoid)
    right_wall = plt.Polygon([
        (outer_bot_w, barrel_bot), (outer_top_w, barrel_top),
        (steps[-1][4], barrel_top), (steps[0][4], barrel_bot)
    ], fc='#8d99ae', ec='#2b2d42', lw=2)
    ax.add_patch(right_wall)
    # Bottom cap (closed side — narrow)
    ax.add_patch(patches.Rectangle((-outer_bot_w, barrel_bot-0.35), 2*outer_bot_w, 0.35,
        fc='#6d7a8a', ec='#2b2d42', lw=2))
    ax.text(0, barrel_bot - 0.18, 'CLOSED', fontsize=7, ha='center', va='center',
        color='white', fontweight='bold')

    # Draw stepped inner wall fill
    for i, (name, typ, y, h, w) in enumerate(steps):
        # Interpolated outer wall position at this height
        frac_bot = (y - barrel_bot) / (barrel_top - barrel_bot)
        frac_top = (y + h - barrel_bot) / (barrel_top - barrel_bot)
        ow_bot = outer_bot_w + (outer_top_w - outer_bot_w) * frac_bot
        ow_top = outer_bot_w + (outer_top_w - outer_bot_w) * frac_top
        # Left fill
        left_trap = plt.Polygon([
            (-ow_bot, y), (-ow_top, y+h), (-w, y+h), (-w, y)
        ], fc='#8d99ae', ec='none', lw=0)
        ax.add_patch(left_trap)
        # Right fill
        right_trap = plt.Polygon([
            (ow_bot, y), (ow_top, y+h), (w, y+h), (w, y)
        ], fc='#8d99ae', ec='none', lw=0)
        ax.add_patch(right_trap)

    # Inner cavity background
    ax.add_patch(patches.Rectangle((-steps[0][4], barrel_bot),
        2*steps[0][4], barrel_top, fc='#edf2f4', ec='none', zorder=0))

    # Draw staircase edges
    ax.plot(left_x, left_y, color='#2b2d42', lw=1.5, zorder=4)
    ax.plot(right_x, right_y, color='#2b2d42', lw=1.5, zorder=4)

    # Draw lenses and spacers
    highlight = 'L3'
    for name, typ, y, h, w in steps:
        is_hl = (name == highlight)
        if typ == 'lens':
            t = np.linspace(-w*0.90, w*0.90, 100)
            curve = 0.22 * np.cos(t * np.pi / (2*w*0.90))
            top_c = y + h*0.65 + curve
            bot_c = y + h*0.15 - 0.10 * np.cos(t * np.pi / (2*w*0.90))
            if is_hl:
                ax.fill_between(t, bot_c-0.06, top_c+0.06, fc='#ffd166', ec='none', alpha=0.5, zorder=5)
                ax.fill_between(t, bot_c, top_c, fc='#06d6a0', ec='#028a5e', lw=2.5, alpha=0.9, zorder=6)
            else:
                ax.fill_between(t, bot_c, top_c, fc='#a8dadc', ec='#457b9d', lw=1.2, alpha=0.8, zorder=5)
            lbl_x = outer_top_w + 1.5
            if is_hl:
                ax.text(lbl_x, y+h/2, f'► {name} (SELECTED)', fontsize=10, fontweight='bold',
                    color='#028a5e', va='center',
                    bbox=dict(boxstyle='round,pad=0.3', fc='#d4f5e9', ec='#028a5e', lw=1.5))
            else:
                ax.text(lbl_x, y+h/2, name, fontsize=9, color='#457b9d', va='center')
        elif typ == 'spacer':
            ax.add_patch(patches.Rectangle((-w*0.90, y+h*0.3), 2*w*0.90, h*0.35,
                fc='#ddd', ec='#999', lw=1, zorder=5))
            ax.text(outer_top_w+1.5, y+h/2, name, fontsize=8, color='#999', va='center')
        else:  # ring/cover
            ax.add_patch(patches.Rectangle((-w*0.90, y+h*0.15), 2*w*0.90, h*0.7,
                fc='#b5838d', ec='#6d4c5e', lw=1.5, zorder=5))
            ax.text(outer_top_w+1.5, y+h/2, name, fontsize=9, fontweight='bold', color='#6d4c5e', va='center')

    # ── DIMENSION ARROWS for L3 (highlighted step) ──
    # Find L3 data
    l3_name, l3_typ, l3_y, l3_h, l3_w = steps[4]  # L3

    # Arrow position X (left of barrel, outside the wall)
    dim_x = -outer_top_w - 1.2

    # Barrel Step Height arrow (full step height)
    ax.annotate('', xy=(dim_x, l3_y), xytext=(dim_x, l3_y + l3_h),
        arrowprops=dict(arrowstyle='<->', color='#e63946', lw=2))
    ax.text(dim_x - 0.3, l3_y + l3_h/2, 'Step\nHeight\n(1.4mm)', fontsize=8,
        fontweight='bold', ha='right', va='center', color='#e63946',
        bbox=dict(boxstyle='round,pad=0.2', fc='#fff5f5', ec='#e63946', lw=1))

    # Lens Height arrow (just the lens portion within the step)
    lens_bot = l3_y + l3_h*0.15
    lens_top = l3_y + l3_h*0.65 + 0.22  # approximate top of lens curve
    lens_dim_x = dim_x + 0.7
    ax.annotate('', xy=(lens_dim_x, lens_bot), xytext=(lens_dim_x, lens_top),
        arrowprops=dict(arrowstyle='<->', color='#028a5e', lw=2))
    ax.text(lens_dim_x + 0.15, (lens_bot + lens_top)/2, 'Lens\nHeight\n(1.085mm)',
        fontsize=8, fontweight='bold', ha='left', va='center', color='#028a5e',
        bbox=dict(boxstyle='round,pad=0.2', fc='#d4f5e9', ec='#028a5e', lw=1))

    # ── GLOBAL DIMENSION ARROWS ──
    # Inner Diameter at bottom (narrowest — closed side)
    ax.annotate('', xy=(-steps[0][4], -1), xytext=(steps[0][4], -1),
        arrowprops=dict(arrowstyle='<->', color='#e63946', lw=2))
    ax.text(0, -1.8, 'Inner Dia (Closed)\n(Ø Narrowest)', fontsize=10, fontweight='bold', ha='center',
        color='#1d3557', bbox=dict(boxstyle='round,pad=0.3', fc='#f1faee', ec='#457b9d'))

    # Inner Diameter at top (widest — open side)
    ax.annotate('', xy=(-steps[-1][4], barrel_top + 0.8), xytext=(steps[-1][4], barrel_top + 0.8),
        arrowprops=dict(arrowstyle='<->', color='#2a9d8f', lw=2))
    ax.text(0, barrel_top + 1.5, 'Inner Dia (Open)\n(Ø Widest)', fontsize=10, fontweight='bold', ha='center',
        color='#264653', bbox=dict(boxstyle='round,pad=0.3', fc='#e9f5db', ec='#2a9d8f'))

    # TTL (total track length)
    ttl_x = -outer_top_w - 3.0
    ax.annotate('', xy=(ttl_x, 0.4), xytext=(ttl_x, 14.3),
        arrowprops=dict(arrowstyle='<->', color='#e63946', lw=2))
    ax.text(ttl_x, 7.5, 'TTL\n8.430mm', fontsize=10, fontweight='bold', ha='center', rotation=90,
        color='#1d3557', bbox=dict(boxstyle='round,pad=0.3', fc='#f1faee', ec='#457b9d'))

    # Open side label at top
    ax.text(0, barrel_top + 3.2, '▼ OPEN SIDE (Insertion from here) ▼', fontsize=11,
        fontweight='bold', ha='center', color='#0077b6',
        bbox=dict(boxstyle='round,pad=0.4', fc='#caf0f8', ec='#0077b6', lw=1.5))

    # Insertion order annotation
    ax.annotate('1st insertion\n(deepest)', xy=(steps[0][4]+0.3, 0.7),
        xytext=(outer_top_w+1.5, -1),
        fontsize=9, fontstyle='italic', color='#e63946', fontweight='bold',
        arrowprops=dict(arrowstyle='->', color='#e63946', lw=1.5, connectionstyle='arc3,rad=-0.2'),
        bbox=dict(boxstyle='round,pad=0.3', fc='#fff5f5', ec='#e63946', alpha=0.9))

    # Click hint
    ax.annotate('Click lens to\nview details →', xy=(l3_w*0.5, l3_y+l3_h/2),
        xytext=(outer_top_w+1.5, 2.5),
        fontsize=10, fontstyle='italic', color='#028a5e', fontweight='bold',
        arrowprops=dict(arrowstyle='->', color='#028a5e', lw=1.5, connectionstyle='arc3,rad=0.3'),
        bbox=dict(boxstyle='round,pad=0.4', fc='#d4f5e9', ec='#028a5e', alpha=0.9))

    ax.set_title('Barrel Assembly — Cross Section (Stepped Interior)\nWide = Open Top, Narrow = Closed Bottom',
        fontsize=14, fontweight='bold', color='#1d3557', pad=15)
    ax.axis('off'); fig.tight_layout()
    fig.savefig(os.path.join(OUT, 'img_barrel.png'), dpi=150, bbox_inches='tight')
    plt.close()

# LENS - Circular 2D top-down view (concentric circles)
def draw_lens():
    fig, ax = plt.subplots(figsize=(10, 8))
    ax.set_xlim(-5.5, 7); ax.set_ylim(-5.5, 5.5); ax.set_aspect('equal')
    
    outer_r = 3.01    # Outer Diameter = 6.020mm -> radius 3.01
    media_r = 2.22    # Media Effective Dia = 4.440mm -> radius 2.22
    light_r = 2.164   # Light Effective Dia = 4.328mm -> radius 2.164
    
    # Draw circles from outer to inner
    ax.add_patch(plt.Circle((0,0), outer_r, fc='#caf0f8', ec='#457b9d', lw=2.5))
    ax.add_patch(plt.Circle((0,0), media_r, fc='#a8dadc', ec='#457b9d', lw=1.5, ls='--'))
    ax.add_patch(plt.Circle((0,0), light_r, fc='#90e0ef', ec='#0077b6', lw=1.5))
    
    # Center dot
    ax.plot(0, 0, '+', color='#e63946', markersize=12, mew=2)
    
    # Gate cutting mark on edge
    theta = np.radians(30)
    gx, gy = outer_r * np.cos(theta), outer_r * np.sin(theta)
    ax.plot([gx-0.15, gx+0.15], [gy-0.1, gy+0.1], color='#e63946', lw=2)
    ax.text(gx+0.3, gy+0.3, 'Gate\nCutting', fontsize=7, color='#e63946', fontstyle='italic')
    
    # Labels on diagram
    ax.text(0, light_r-0.4, 'Light Eff.\nDia', fontsize=7, ha='center', color='#0077b6', alpha=0.7)
    ax.text(0, -media_r+0.3, 'Media Eff. Dia', fontsize=7, ha='center', color='#457b9d', alpha=0.7)
    
    # Param 1: Outer Diameter
    ax.annotate('', xy=(-outer_r, -4.3), xytext=(outer_r, -4.3),
        arrowprops=dict(arrowstyle='<->', color='#e63946', lw=2))
    ax.text(0, -4.9, 'Outer Diameter (Ø6.020 mm)', fontsize=11, fontweight='bold', ha='center',
        color='#1d3557', bbox=dict(boxstyle='round,pad=0.3', fc='#f1faee', ec='#457b9d'))
    
    # Param 2: Light Effective Diameter
    ax.annotate('', xy=(-light_r, 4), xytext=(light_r, 4),
        arrowprops=dict(arrowstyle='<->', color='#2a9d8f', lw=2))
    ax.text(0, 4.6, 'Light Effective Dia (Ø4.328 mm)', fontsize=10, fontweight='bold', ha='center',
        color='#264653', bbox=dict(boxstyle='round,pad=0.3', fc='#e9f5db', ec='#2a9d8f'))
    
    # Param 3: Center Thickness (shown as text annotation since top-down)
    ax.annotate('Center Thickness\n= 1.0850 mm\n(side dimension)', xy=(0, 0), xytext=(4.5, 3),
        fontsize=10, fontweight='bold', color='#1d3557',
        arrowprops=dict(arrowstyle='->', color='#e63946', lw=1.5, connectionstyle='arc3,rad=0.2'),
        bbox=dict(boxstyle='round,pad=0.4', fc='#f1faee', ec='#457b9d'))
    
    ax.set_title('Lens — Top View (e.g. L3: APEL5514ML)', fontsize=14,
        fontweight='bold', color='#1d3557', pad=15)
    ax.axis('off'); fig.tight_layout()
    fig.savefig(os.path.join(OUT, 'img_lens.png'), dpi=150, bbox_inches='tight')
    plt.close()

# SPACER - Circular 2D ring top-down view
def draw_spacer():
    fig, ax = plt.subplots(figsize=(10, 8))
    ax.set_xlim(-5.5, 7); ax.set_ylim(-5.5, 5.5); ax.set_aspect('equal')
    
    outer_r = 2.9   # Outer Dia = 5.800mm
    inner_r = 2.2   # Inner Dia = 4.400mm
    
    # Outer circle (spacer body)
    ax.add_patch(plt.Circle((0,0), outer_r, fc='#d5c4a1', ec='#6b4226', lw=2.5))
    # Inner circle (hole - white)
    ax.add_patch(plt.Circle((0,0), inner_r, fc='#ffffff', ec='#6b4226', lw=2, ls='--'))
    
    # Center label
    ax.text(0, 0, 'Open\nCenter', fontsize=9, ha='center', va='center',
        color='#999', fontstyle='italic')
    ax.plot(0, 0, '+', color='#ccc', markersize=10, mew=1)
    
    # Param 1: Outer Diameter
    ax.annotate('', xy=(-outer_r, -4.3), xytext=(outer_r, -4.3),
        arrowprops=dict(arrowstyle='<->', color='#e63946', lw=2))
    ax.text(0, -4.9, 'Outer Diameter (Ø5.800 mm)', fontsize=11, fontweight='bold', ha='center',
        color='#1d3557', bbox=dict(boxstyle='round,pad=0.3', fc='#f1faee', ec='#457b9d'))
    
    # Param 2: Inner Diameter
    ax.annotate('', xy=(-inner_r, 4), xytext=(inner_r, 4),
        arrowprops=dict(arrowstyle='<->', color='#2a9d8f', lw=2))
    ax.text(0, 4.6, 'Inner Diameter (Ø4.400 mm)', fontsize=10, fontweight='bold', ha='center',
        color='#264653', bbox=dict(boxstyle='round,pad=0.3', fc='#e9f5db', ec='#2a9d8f'))
    
    # Param 3: Thickness (side dimension note)
    ax.annotate('Thickness\n= 0.235 mm\n(side dimension)', xy=(outer_r*0.7, outer_r*0.7),
        xytext=(4.5, 3),
        fontsize=10, fontweight='bold', color='#1d3557',
        arrowprops=dict(arrowstyle='->', color='#e63946', lw=1.5, connectionstyle='arc3,rad=0.2'),
        bbox=dict(boxstyle='round,pad=0.4', fc='#f1faee', ec='#457b9d'))
    
    ax.set_title('Spacer Ring — Top View (FILM type)', fontsize=14,
        fontweight='bold', color='#1d3557', pad=15)
    ax.axis('off'); fig.tight_layout()
    fig.savefig(os.path.join(OUT, 'img_spacer.png'), dpi=150, bbox_inches='tight')
    plt.close()

print("Regenerating 3 images...")
draw_barrel(); print("  [1/3] Barrel (cone) done")
draw_lens(); print("  [2/3] Lens (2D circular) done")
draw_spacer(); print("  [3/3] Spacer (2D ring) done")

# Rebuild Word doc
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'; style.font.size = Pt(11)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for _ in range(6): p.add_run('\n')
r = p.add_run('Model Editor\n'); r.font.size = Pt(36); r.bold = True; r.font.color.rgb = RGBColor(0x1d,0x35,0x57)
r = p.add_run('UI Design Document\n'); r.font.size = Pt(24); r.font.color.rgb = RGBColor(0x45,0x7b,0x9d)
r = p.add_run('\nCAD Base Parameters → Visual Editor\n'); r.font.size = Pt(14); r.font.color.rgb = RGBColor(0x66,0x66,0x66)

pages = [
    ('img_lens_tray.png', 'Lens Tray — Top View',
     'Each machine has its own lens tray containing one type of lens.\n'
     'Base Parameters:\n• Tray Hole Count (X × Y)\n• Tray Hole Interval (X, Y) mm\n• 1st Hole Start Position (X, Y) mm', 6),
    ('img_barrel_tray.png', 'Barrel Tray (Assy Tray) — Top View',
     'Barrel tray holds multiple barrels in a grid on the conveyor belt.\n'
     'Base Parameters:\n• Tray Hole Count (X × Y)\n• Tray Hole Interval (X, Y) mm\n• 1st Hole Start Position (X, Y) mm', 6),
    ('img_barrel.png', 'Barrel Assembly — Cross Section (Stepped Interior)',
     'Wide open side at top (insertion from here), narrow closed side at bottom.\n'
     'Stepped/staircase interior with each ledge wider than the one below.\n'
     'L3 highlighted with Step Height and Lens Height dimension arrows.\n'
     'Base Parameters:\n• Inner Diameter (Closed/Open)\n• TTL\n• Step Height per position\n• Lens/Spacer Height', 5.0),
    ('img_lens.png', 'Lens — Top View (Opened from Barrel)',
     'Circular 2D view of the selected lens showing concentric diameter zones.\n'
     'This view opens when user clicks a lens in the barrel diagram.\n'
     'Base Parameters:\n• Outer Diameter\n• Light Effective Diameter\n• Center Thickness', 5.5),
    ('img_spacer.png', 'Spacer Ring — Top View',
     'Ring-shaped spacer viewed from top. Each spacer varies per barrel position.\n'
     'Base Parameters:\n• Outer Diameter\n• Inner Diameter\n• Thickness', 5.5),
]

for img_file, title, desc, w in pages:
    doc.add_page_break()
    h = doc.add_heading(title, level=1)
    for run in h.runs: run.font.color.rgb = RGBColor(0x1d,0x35,0x57)
    doc.add_picture(os.path.join(OUT, img_file), width=Inches(w))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('')
    doc.add_paragraph(desc)

out_path = os.path.join(OUT, 'Model_Editor_UI_Design.docx')
doc.save(out_path)
print(f"\nWord document saved: {out_path}")
