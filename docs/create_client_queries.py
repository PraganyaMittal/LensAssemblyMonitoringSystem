import matplotlib
matplotlib.use('Agg')
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

OUT = os.path.dirname(os.path.abspath(__file__))

doc = Document()

# ── Page margins ──
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Style setup ──
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(6)

# Heading styles
for level in [1, 2, 3]:
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Calibri'
    h.font.color.rgb = RGBColor(0x1d, 0x35, 0x57)

# ═══════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for _ in range(5):
    p.add_run('\n')
r = p.add_run('Model Editor\n')
r.font.size = Pt(36)
r.bold = True
r.font.color.rgb = RGBColor(0x1d, 0x35, 0x57)
r = p.add_run('Client Queries Document\n')
r.font.size = Pt(24)
r.font.color.rgb = RGBColor(0x45, 0x7b, 0x9d)
r = p.add_run('\nImplementation Blockers & Clarification Requests\n')
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
r = p.add_run('\n\n\n')
r = p.add_run('Date: _______________\n')
r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

# ═══════════════════════════════════════════════════════════════
# CONTEXT PAGE
# ═══════════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading('Context', level=1)

p = doc.add_paragraph()
p.add_run('What we are building: ').bold = True
p.add_run(
    'A web-based Model Editor that allows engineers to enter CAD base parameters, '
    'auto-derive SpecData values using formulas, preview changes via diff, '
    'and deploy models to machines.'
)

p = doc.add_paragraph()
p.add_run('What we already know: ').bold = True

bullets = [
    'Factory layout: Conveyor belt with 5-10 machines per line, each machine inserts one lens + one spacer',
    'Tray parameters (Hole Count, Interval, 1st Hole Position) — these are known and exist directly in SpecData XML',
    'Picker mechanism parameters (Z-positions, Pressure, Angle) — structure is known',
    'XML format: data/group/spec/val hierarchy — same as existing system',
    'Model structure: Same for all machines in a line, but each machine has different key parameter values',
    'Assembly sequence: Configurable per line (not fixed)',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

p = doc.add_paragraph()
p.add_run('What we need from you (this document): ').bold = True
p.add_run(
    'The following questions are implementation blockers. '
    'Without these answers, we cannot proceed with the derivation engine and complete parameter mapping. '
    'Please answer each question or provide the referenced documents/data.'
)

# ═══════════════════════════════════════════════════════════════
# Helper: Add a query table
# ═══════════════════════════════════════════════════════════════
query_num = [0]

def add_query(question, why_needed, answer_format="", reference_image=None):
    query_num[0] += 1
    qid = f"Q{query_num[0]}"

    # Query heading
    p = doc.add_paragraph()
    r = p.add_run(f'{qid}. ')
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x0e, 0xa5, 0xe9)
    r = p.add_run(question)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x1d, 0x35, 0x57)

    # Why we need this
    p = doc.add_paragraph()
    r = p.add_run('Why we need this: ')
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0xe6, 0x39, 0x46)
    r = p.add_run(why_needed)
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Expected answer format
    if answer_format:
        p = doc.add_paragraph()
        r = p.add_run('Expected answer format: ')
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0x2a, 0x9d, 0x8f)
        r = p.add_run(answer_format)
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Reference image
    if reference_image:
        img_path = os.path.join(OUT, reference_image)
        if os.path.exists(img_path):
            p = doc.add_paragraph()
            r = p.add_run('Reference: ')
            r.bold = True
            r.font.size = Pt(10)
            doc.add_picture(img_path, width=Inches(4.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Answer box
    answer_table = doc.add_table(rows=1, cols=1)
    answer_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = answer_table.cell(0, 0)
    cell.text = 'Client Answer:\n\n\n\n'
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    cell.paragraphs[0].runs[0].font.italic = True
    # Cell shading
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): 'F8F9FA'
    })
    shading.append(shading_elm)

    # Separator
    doc.add_paragraph('─' * 60).runs[0].font.color.rgb = RGBColor(0xDD, 0xDD, 0xDD)


# ═══════════════════════════════════════════════════════════════
# SECTION 1: DERIVATION FORMULAS (Critical)
# ═══════════════════════════════════════════════════════════════
doc.add_page_break()
h = doc.add_heading('Section 1: Derivation Formulas', level=1)
p = doc.add_paragraph()
r = p.add_run('⚠ CRITICAL — These are the #1 implementation blocker')
r.bold = True
r.font.color.rgb = RGBColor(0xe6, 0x39, 0x46)
r.font.size = Pt(12)

add_query(
    'What are the exact formulas to convert CAD base parameters to SpecData values?',
    'The Model Editor needs to auto-calculate derived values when the engineer enters CAD base parameters. '
    'Without the formulas, we can only build a manual editor (no automation).',
    'For each derived parameter, please provide:\n'
    'derived_param_name = formula(base_param_1, base_param_2, ...)\n'
    'Example: spec_z_position = SAG1 + center_thickness × 0.5',
    'img_lens.png'
)

add_query(
    'Are the derivation formulas the same for all lens positions (L1 through L7), '
    'or does each position have different formulas?',
    'If formulas differ per position, we need separate formula sets for each machine in the line.',
    'Same for all / Different per position (please specify which differ)'
)

add_query(
    'Are there lookup tables or conditional rules involved in the derivation, '
    'or are all derivations purely mathematical?',
    'Lookup tables and conditional logic require different implementation than simple math formulas. '
    'We need to know the complexity to design the formula engine correctly.',
    'Pure math / Lookup tables (please provide tables) / Conditional rules (please describe)'
)

add_query(
    'Can you provide a worked example? Given a specific CAD drawing, '
    'show us: here are the base params → here are the derived SpecData values.',
    'A worked example lets us verify our formula implementation is correct. '
    'Even one complete example for one lens position would be very helpful.',
    'CAD base param values → calculated intermediate values → final SpecData values'
)

# ═══════════════════════════════════════════════════════════════
# SECTION 2: COMPLETE PARAMETER LIST
# ═══════════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading('Section 2: Complete Parameter Lists', level=1)

add_query(
    'What is the complete list of CAD base parameters for a LENS?',
    'We need every base parameter that the engineer reads from the CAD drawing and enters into the system. '
    'We currently know some from the CAD (SAG, thickness, diameters) but need the complete list.',
    'Parameter Name | Unit | Example Value | Description\n'
    '(We have identified these so far: S1 SAG, S2 SAG, Center Thickness, '
    'Light Effective Dia, Media Effective Dia, Outer Diameter, Refractive Index, Dispersion, Gate Cutting)',
    'img_lens.png'
)

add_query(
    'What is the complete list of CAD base parameters for a SPACER?',
    'Same as above but for spacer components. We currently know: Outer Diameter, Inner Diameter, Thickness, Material.',
    'Parameter Name | Unit | Example Value | Description',
    'img_spacer.png'
)

add_query(
    'What is the complete list of CAD base parameters for the BARREL?',
    'Barrel dimensions that feed into SpecData derivation. '
    'We currently know: Inner Diameter, Outer Diameter, TTL.',
    'Parameter Name | Unit | Example Value | Description',
    'img_barrel.png'
)

add_query(
    'What is the complete list of DERIVED SpecData parameters that change between models?',
    'We need to know exactly which SpecData XML parameters are modified when creating a new model from DEFAULT_MODEL. '
    'This tells us the output of the derivation formulas.',
    'SpecData Parameter Name (as it appears in XML) | Which base params it depends on | Formula'
)

# ═══════════════════════════════════════════════════════════════
# SECTION 3: VALIDATION & LIMITS
# ═══════════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading('Section 3: Validation & Limits', level=1)

add_query(
    'Which parameters have validation limits (min/max/range)?',
    'We want to prevent invalid values from being entered. '
    'For example: Z positions must be positive, pressure must be within safe range, etc.',
    'Parameter Name | Min Value | Max Value | Unit'
)

add_query(
    'Are there dependencies between parameters? '
    '(e.g., changing lens diameter affects picker approach position)',
    'If parameters are interdependent, we need to implement cascading updates '
    'so that changing one value automatically adjusts related values.',
    'Parameter A | affects → | Parameter B | How (formula/rule)'
)

add_query(
    'For the Spacer Picker — does it have Push Pos Z and Push Ready Pos Z parameters, '
    'or are these only for the Lens Picker?',
    'The Lens Picker has 6 parameters (including Push positions), '
    'but the Spacer Picker was listed with only 4. We need to confirm if this is correct or incomplete.',
    'Spacer Picker has Push params: Yes / No'
)

# ═══════════════════════════════════════════════════════════════
# SECTION 4: WORKFLOW CONFIRMATION
# ═══════════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading('Section 4: Workflow Confirmation', level=1)

p = doc.add_paragraph()
p.add_run('Please confirm or correct the proposed workflow:').bold = True

doc.add_paragraph(
    'We propose the following workflow for creating/editing a model through the web application:'
)
steps = [
    'Engineer opens Model Editor for a specific machine (e.g., MC-3 which handles L3 + SP3)',
    'System loads DEFAULT_MODEL as the base template',
    'Engineer sees visual diagrams of components (barrel cross-section, lens, spacer, trays, picker)',
    'Engineer enters CAD base parameters into input fields mapped to diagram annotations',
    'System auto-calculates derived SpecData values using formulas',
    'Engineer reviews diff (before vs after comparison of all changed parameters)',
    'Engineer deploys the updated model to the machine',
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f'Step {i}: {step}', style='List Number')

add_query(
    'Is this proposed workflow correct? Any steps missing or in wrong order?',
    'We need workflow confirmation before building the UI. '
    'Any correction now saves significant development time later.',
    'Correct / Incorrect (please describe the correct flow)'
)

add_query(
    'Should users be able to manually override auto-derived values? '
    'We propose: Yes, with a warning indicator.',
    'Some engineers may need to fine-tune derived values. '
    'We want to allow this but show a clear warning that the value was manually overridden.',
    'Allow override with warning / Allow override without warning / Do not allow override'
)

add_query(
    'Is the per-machine editor approach correct? '
    'Or should all machines in a line be editable from one screen?',
    'Currently we designed the editor to show one machine at a time '
    '(only its lens + spacer params), with the full barrel as context. '
    'Please confirm this is the right approach.',
    'Per-machine (current design) / All machines on one screen'
)

# ═══════════════════════════════════════════════════════════════
# SECTION 5: UI DESIGN CONFIRMATION
# ═══════════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading('Section 5: UI Design Confirmation', level=1)

p = doc.add_paragraph()
p.add_run('Please review the following diagrams from our UI Design Document ').bold = True
p.add_run('(attached separately) and confirm or suggest changes:')

add_query(
    'Are the component diagrams sufficient? Are any diagrams missing?',
    'We have created diagrams for: Lens Tray, Barrel Tray, Barrel Cross-Section, '
    'Lens (top view), Spacer (top view). We want to confirm nothing is missing.',
    'Sufficient / Missing: (please list what is missing)',
    'img_barrel.png'
)

add_query(
    'Should the barrel cross-section show exact physical proportions from CAD, '
    'or is a schematic representation sufficient?',
    'Exact proportions require precise CAD data. Schematic is simpler but less accurate.',
    'Exact proportions / Schematic is OK'
)

# ═══════════════════════════════════════════════════════════════
# SECTION 6: DATA REQUEST
# ═══════════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading('Section 6: Data & Document Requests', level=1)

p = doc.add_paragraph()
p.add_run('To proceed with implementation, we also need the following data/documents:').bold = True

data_requests = [
    ('Sample SpecData XML file',
     'A complete SpecData XML file from DEFAULT_MODEL (Recipe.xml from the View folder). '
     'This is critical for us to map parameters to the correct XML tags and attributes.'),
    ('Complete Model folder structure',
     'Full listing of all files and folders inside a model directory '
     '(Algo, View, and any other sibling folders with their file names).'),
    ('Sample CAD drawing with corresponding SpecData',
     'One complete CAD drawing (any product) together with the SpecData values '
     'that were derived from it. This serves as our verification reference.'),
]

for title, desc in data_requests:
    p = doc.add_paragraph()
    r = p.add_run(f'📋 {title}')
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x1d, 0x35, 0x57)
    p = doc.add_paragraph(desc)
    p.paragraph_format.space_after = Pt(12)

# ═══════════════════════════════════════════════════════════════
# SUMMARY TABLE
# ═══════════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading('Summary of All Queries', level=1)

summary_table = doc.add_table(rows=1, cols=4)
summary_table.style = 'Light Grid Accent 1'
summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header row
headers = ['#', 'Question Summary', 'Priority', 'Status']
for i, header in enumerate(headers):
    cell = summary_table.rows[0].cells[i]
    cell.text = header
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(9)

# Data rows
queries_summary = [
    ('Q1', 'Derivation formulas (CAD → SpecData)', '🔴 Critical'),
    ('Q2', 'Formulas same for all positions?', '🔴 Critical'),
    ('Q3', 'Lookup tables or conditional rules?', '🔴 Critical'),
    ('Q4', 'Worked example (CAD → derived values)', '🔴 Critical'),
    ('Q5', 'Complete lens base param list', '🟡 High'),
    ('Q6', 'Complete spacer base param list', '🟡 High'),
    ('Q7', 'Complete barrel base param list', '🟡 High'),
    ('Q8', 'Complete derived SpecData param list', '🟡 High'),
    ('Q9', 'Validation limits per parameter', '🟢 Medium'),
    ('Q10', 'Parameter dependencies', '🟢 Medium'),
    ('Q11', 'Spacer picker push params', '🟢 Medium'),
    ('Q12', 'Workflow confirmation', '🟡 High'),
    ('Q13', 'Override policy confirmation', '🟢 Medium'),
    ('Q14', 'Per-machine editor confirmation', '🟡 High'),
    ('Q15', 'Diagram completeness check', '🟢 Medium'),
    ('Q16', 'Proportions: exact vs schematic', '🟢 Medium'),
]

for qid, summary, priority in queries_summary:
    row = summary_table.add_row()
    row.cells[0].text = qid
    row.cells[1].text = summary
    row.cells[2].text = priority
    row.cells[3].text = 'Pending'
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)

# Set column widths
for row in summary_table.rows:
    row.cells[0].width = Cm(1.5)
    row.cells[1].width = Cm(9)
    row.cells[2].width = Cm(2.5)
    row.cells[3].width = Cm(2)

# ═══════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════
out_path = os.path.join(OUT, 'Client_Queries.docx')
doc.save(out_path)
print(f"Client Queries document saved: {out_path}")
print(f"Total queries: {query_num[0]}")
