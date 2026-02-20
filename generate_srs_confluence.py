from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

doc = Document()

# Styles setup
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(4)

for level in range(1, 5):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Calibri'
    if level == 1:
        hs.font.size = Pt(18)
        hs.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    elif level == 2:
        hs.font.size = Pt(15)
        hs.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    elif level == 3:
        hs.font.size = Pt(13)
        hs.font.color.rgb = RGBColor(0x37, 0x37, 0x37)
    elif level == 4:
        hs.font.size = Pt(11)
        hs.font.color.rgb = RGBColor(0x37, 0x37, 0x37)

def add_horizontal_line(doc):
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_table(doc, rows_data):
    table = doc.add_table(rows=len(rows_data), cols=len(rows_data[0]))
    table.style = 'Light Grid Accent 1'
    for i, row_data in enumerate(rows_data):
        for j, cell_text in enumerate(row_data):
            cell = table.cell(i, j)
            cell.text = cell_text.strip()
            for p in cell.paragraphs:
                p.style.font.size = Pt(10)
    doc.add_paragraph()

def process_text_with_bold_and_code(paragraph, text):
    """Handle **bold**, `code`, and plain text."""
    parts = re.split(r'(\*\*.*?\*\*|`[^`]+`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        else:
            paragraph.add_run(part)

# Read the markdown file
with open(r'C:\Users\Divyansh V\.gemini\antigravity\brain\ff22f578-df62-46e2-a25e-b811752a40a8\srs_confluence_entries.md', 'r', encoding='utf-8') as f:
    content = f.read()

# Title page
title = doc.add_heading('Factory Monitoring System', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_heading('SRS Entries for Confluence', level=1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('Date: February 19, 2026\n').bold = True
meta.add_run('Version: 1.0\n')
meta.add_run('Team: Divyansh, Bhavik, Praganya\n')
doc.add_page_break()

# Parse and write each SRS section
srs_sections = content.split('\n---\n')

for section in srs_sections:
    lines = section.strip().split('\n')
    if not lines:
        continue
    
    in_table = False
    table_rows = []
    skip_note_block = False
    
    for i, line in enumerate(lines):
        stripped = line.strip()
        
        # Skip the top-level title and note block
        if stripped.startswith('# Factory Monitoring') or stripped.startswith('All 12 requirements'):
            continue
        if stripped.startswith('> [!NOTE]') or stripped.startswith('> **Changes'):
            skip_note_block = True
            continue
        if skip_note_block:
            if stripped.startswith('> -') or stripped.startswith('>'):
                continue
            else:
                skip_note_block = False
                if not stripped:
                    continue
        
        # Empty line
        if not stripped:
            if in_table and table_rows:
                add_table(doc, table_rows)
                table_rows = []
                in_table = False
            continue
        
        # Table rows
        if stripped.startswith('|') and stripped.endswith('|'):
            cells = [c.strip() for c in stripped.split('|')[1:-1]]
            # Skip separator rows like |---|---|
            if all(set(c) <= set('-: ') for c in cells):
                continue
            in_table = True
            table_rows.append(cells)
            continue
        
        if in_table and table_rows:
            add_table(doc, table_rows)
            table_rows = []
            in_table = False
        
        # Headings
        if stripped.startswith('## SRS-') or stripped.startswith('## NFR-'):
            doc.add_heading(stripped[3:], level=1)
            continue
        if stripped.startswith('### '):
            doc.add_heading(stripped[4:], level=2)
            continue
        if stripped.startswith('#### '):
            doc.add_heading(stripped[5:], level=3)
            continue
        
        # Type/Priority line
        if stripped.startswith('**Type:**'):
            p = doc.add_paragraph()
            clean = stripped.replace('&nbsp;', ' ').replace('|', '|')
            process_text_with_bold_and_code(p, clean)
            continue
        
        # Bullet points
        if stripped.startswith('- **') or stripped.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            text = stripped[2:]
            process_text_with_bold_and_code(p, text)
            continue
        
        # Sub-bullets
        if stripped.startswith('  - '):
            p = doc.add_paragraph(style='List Bullet 2')
            text = stripped[4:]
            process_text_with_bold_and_code(p, text)
            continue
        
        # Regular paragraph
        p = doc.add_paragraph()
        process_text_with_bold_and_code(p, stripped)
    
    # Flush any remaining table
    if in_table and table_rows:
        add_table(doc, table_rows)
    
    # Add separator between SRS entries
    add_horizontal_line(doc)

output_path = r'c:\Users\Divyansh V\Desktop\Factory Monitor OK\SRS_Confluence_Entries.docx'
doc.save(output_path)
print(f"Done! Saved to: {output_path}")
