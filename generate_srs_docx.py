"""
Generate SRS_Document.docx — IEEE 830-style SRS
Includes: system requirements, hardware/software specs, hosting, deployment, and project overview.
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

OUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "SRS_FactoryMonitoring.docx")

def shade(cell, color):
    cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>'))

def style_tbl(t, hdr_bg="1F4E79", hdr_fg="FFFFFF"):
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for c in t.rows[0].cells:
        shade(c, hdr_bg)
        for p in c.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor.from_string(hdr_fg)
                r.font.bold = True; r.font.size = Pt(9)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
    for i, row in enumerate(t.rows[1:], 1):
        bg = "F2F7FB" if i % 2 == 0 else "FFFFFF"
        for c in row.cells:
            shade(c, bg)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.space_before = Pt(2)

def tbl(doc, hdrs, rows, widths=None, bg="1F4E79"):
    t = doc.add_table(rows=1+len(rows), cols=len(hdrs))
    t.style = 'Table Grid'
    for j, h in enumerate(hdrs):
        t.rows[0].cells[j].text = h
    for i, rd in enumerate(rows):
        for j, v in enumerate(rd):
            t.rows[i+1].cells[j].text = v
    if widths:
        for row in t.rows:
            for j, w in enumerate(widths):
                row.cells[j].width = Cm(w)
    style_tbl(t, hdr_bg=bg)
    return t

def build():
    doc = Document()

    # Page setup
    s = doc.sections[0]
    s.page_width = Inches(8.5); s.page_height = Inches(11)
    s.top_margin = Cm(2.5); s.bottom_margin = Cm(2.5)
    s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'; style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    for lv in range(1, 4):
        hs = doc.styles[f'Heading {lv}']
        hs.font.name = 'Calibri'
        hs.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        hs.font.bold = True
        hs.font.size = Pt([0, 22, 16, 13][lv])

    # ── COVER PAGE ──
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Software Requirements Specification (SRS)")
    r.font.size = Pt(28); r.font.bold = True; r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Factory Monitoring System")
    r.font.size = Pt(20); r.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

    doc.add_paragraph()
    for label, val in [("Version:", "1.0"), ("Date:", "February 13, 2026"),
                       ("Project Duration:", "December 2025 – Present"),
                       ("Prepared By:", "Divyansh, Bhavik, Praganya")]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rl = p.add_run(label + " "); rl.font.bold = True; rl.font.size = Pt(12)
        rl.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        rv = p.add_run(val); rv.font.size = Pt(12); rv.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    doc.add_page_break()

    # ── TABLE OF CONTENTS ──
    doc.add_heading("Table of Contents", level=1)
    toc = [
        "1. Introduction",
        "   1.1 Purpose", "   1.2 Scope", "   1.3 Intended Audience",
        "   1.4 Project Team", "   1.5 Definitions & Abbreviations", "   1.6 References",
        "2. Overall Description",
        "   2.1 Product Perspective", "   2.2 Product Functions", "   2.3 User Classes",
        "   2.4 Constraints & Assumptions",
        "3. System Requirements",
        "   3.1 Server (Hosting) PC Requirements", "   3.2 Factory Agent PC Requirements",
        "   3.3 Client (Dashboard) PC Requirements", "   3.4 Network Requirements",
        "   3.5 Software Dependencies — Server", "   3.6 Software Dependencies — Frontend",
        "   3.7 Software Dependencies — Agent", "   3.8 Database Requirements",
        "4. System Architecture",
        "   4.1 Architecture Diagram", "   4.2 Component Overview",
        "   4.3 Communication Protocols", "   4.4 Data Flow",
        "5. Functional Requirements",
        "   5.1 Real-Time Monitoring", "   5.2 PC Management",
        "   5.3 Model Library & Distribution", "   5.4 Model Editor & Version History",
        "   5.5 Log Analyser", "   5.6 Agent Communication",
        "6. Non-Functional Requirements",
        "7. Database Schema Summary",
        "8. API Endpoint Summary",
        "9. Deployment & Installation Guide",
        "10. Glossary",
    ]
    for item in toc:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
        p.runs[0].font.size = Pt(11)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════
    # 1. INTRODUCTION
    # ═══════════════════════════════════════════════════
    doc.add_heading("1. Introduction", level=1)

    doc.add_heading("1.1 Purpose", level=2)
    doc.add_paragraph(
        "This Software Requirements Specification (SRS) document describes the complete requirements "
        "for the Factory Monitoring System — a full-stack, enterprise-grade solution for remotely "
        "monitoring factory production-line PCs (Machine Controllers), managing AI/ML model deployment, "
        "distributing configuration files, and analyzing operational logs. This document is intended "
        "for the Project Lead and serves as the definitive specification for the system."
    )

    doc.add_heading("1.2 Scope", level=2)
    doc.add_paragraph(
        "The Factory Monitoring System is a 4-component distributed application deployed on a factory "
        "Local Area Network (LAN). It consists of:"
    )
    for item in [
        "A C++ Agent application running on each factory PC (Machine Controller).",
        "An ASP.NET Core Web API backend hosted on a central server.",
        "A React-based web dashboard accessed via browser by factory engineers.",
        "A SQL Server database for persistent data storage.",
    ]:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading("1.3 Intended Audience", level=2)
    tbl(doc, ["Audience", "Usage"],
        [["Project Lead", "Milestone tracking, feature sign-off, resource planning"],
         ["Developers", "Implementation reference, system specs, dependency management"],
         ["IT / DevOps", "Deployment, hosting requirements, infrastructure planning"],
         ["QA / Testers", "Functional and non-functional test cases"]],
        widths=[4, 13])

    doc.add_heading("1.4 Project Team", level=2)
    tbl(doc, ["Name", "Role", "Modules Owned"],
        [["Divyansh", "Intern — Model Management & UI",
          "Dashboard, model distribution, remote config, UI/UX, sidebar, theme"],
         ["Bhavik", "Intern — Log Analyser",
          "Log parsing, Gantt charts, barrel execution, inspection images"],
         ["Praganya", "Intern — Model Library & UI",
          "Model Library, upload/download, versioning, diff viewer, XML editor"]],
        widths=[3, 5, 9])

    doc.add_heading("1.5 Definitions & Abbreviations", level=2)
    tbl(doc, ["Term", "Definition"],
        [["MC", "Machine Controller — A factory PC on a production line"],
         ["Line", "A production line containing multiple MCs"],
         ["Model Version", "Software/hardware version of PCs (e.g., 3.5, 4.0)"],
         ["Model (AI)", "Machine learning model package (.zip) distributed to MCs"],
         ["Agent", "C++ application on each MC for server communication"],
         ["Heartbeat", "Periodic status signal from Agent to Server"],
         ["Target Model", "Designated standard model for a production line"],
         ["NG", "No Good — Factory quality term for failed inspection"],
         ["SRS", "Software Requirements Specification"],
         ["LAN", "Local Area Network"]],
        widths=[4, 13])

    doc.add_heading("1.6 References", level=2)
    for r in ["IEEE 830-1998 Standard for SRS", "Project README.md",
              "Database Scripts: db/01_CreateDatabase.sql – db/06_UpdateSampleData.sql",
              "Microsoft .NET 8 Documentation", "React 18 Documentation"]:
        doc.add_paragraph(r, style='List Bullet')

    doc.add_page_break()

    # ═══════════════════════════════════════════════════
    # 2. OVERALL DESCRIPTION
    # ═══════════════════════════════════════════════════
    doc.add_heading("2. Overall Description", level=1)

    doc.add_heading("2.1 Product Perspective", level=2)
    doc.add_paragraph(
        "The Factory Monitoring System is a standalone, custom-built enterprise application deployed "
        "entirely within a factory's private LAN. It replaces manual model deployment (via USB drives), "
        "manual configuration editing (via remote desktop), and manual log analysis (via file copying) "
        "with a centralized, web-based, automated solution."
    )

    doc.add_heading("2.2 Product Functions", level=2)
    funcs = [
        ("Real-Time PC Monitoring", "Live dashboard showing online/offline status, application heartbeat, model compliance for all factory PCs grouped by production line."),
        ("Remote Configuration", "Upload, download, and push changes to PC configuration files (.ini, .json) remotely through the web interface."),
        ("Model Library", "Central repository for ML model packages (.zip) with metadata, search, categorization."),
        ("Model Distribution", "Deploy models to specific PCs, entire lines, specific versions, or all PCs — with compliance tracking."),
        ("Model Editor", "In-browser code editor for viewing/editing files inside model archives with syntax highlighting, diff view, undo/redo."),
        ("Version History", "Track all model edits with version snapshots, side-by-side diff comparison, and rollback capability."),
        ("Log Analytics", "Parse factory operation logs into interactive Gantt chart timelines, barrel execution breakdowns, and inspection image galleries."),
        ("Agent Communication", "Automated agent registration, heartbeat, model/config sync, log/image upload, and remote command execution."),
    ]
    for name, desc in funcs:
        p = doc.add_paragraph()
        r = p.add_run(f"{name}: "); r.font.bold = True
        p.add_run(desc)

    doc.add_heading("2.3 User Classes", level=2)
    tbl(doc, ["User Class", "Description", "Access Level"],
        [["Factory Operator", "Monitors dashboard, views PC status", "Read-only dashboard"],
         ["Factory Engineer", "Uploads models, edits configs, analyzes logs", "Full read/write"],
         ["Administrator", "Manages PCs, lines, system settings", "Full admin"]],
        widths=[4, 8, 5])

    doc.add_heading("2.4 Constraints & Assumptions", level=2)
    doc.add_paragraph("Constraints:", style='List Bullet')
    for c in [
        "All components must run on Windows (Agent uses Win32 API).",
        "Database storage only — no network file shares required.",
        "System operates on isolated factory LAN (no internet required after initial setup).",
        "No authentication/authorization currently implemented (trusted network assumption).",
    ]:
        p = doc.add_paragraph(c)
        p.paragraph_format.left_indent = Cm(1.5)
    doc.add_paragraph("Assumptions:", style='List Bullet')
    for a in [
        "All factory PCs are on the same LAN subnet or can reach the server's IP.",
        "Factory PCs run Windows 10 or later.",
        "Server PC has sufficient storage for model archives and log caches.",
    ]:
        p = doc.add_paragraph(a)
        p.paragraph_format.left_indent = Cm(1.5)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════
    # 3. SYSTEM REQUIREMENTS ← THE KEY SECTION
    # ═══════════════════════════════════════════════════
    doc.add_heading("3. System Requirements", level=1)
    doc.add_paragraph(
        "This section specifies the hardware and software requirements for each component of the system."
    )

    # 3.1 Server
    doc.add_heading("3.1 Server (Hosting) PC Requirements", level=2)
    doc.add_paragraph("The server PC hosts both the ASP.NET Core Web API and the SQL Server database.")

    doc.add_heading("Hardware — Minimum", level=3)
    tbl(doc, ["Component", "Minimum", "Recommended"],
        [["Processor", "Intel Core i5 (8th Gen) or equivalent", "Intel Core i7 / Xeon (10th Gen+)"],
         ["RAM", "8 GB", "16 GB or more"],
         ["Storage", "50 GB free (SSD preferred)", "100 GB+ SSD (NVMe recommended)"],
         ["Network", "1 Gbps Ethernet (LAN)", "1 Gbps Ethernet (LAN)"],
         ["Display", "Not required (headless OK)", "Optional for monitoring"]],
        widths=[3, 6.5, 7.5])

    doc.add_heading("Operating System", level=3)
    tbl(doc, ["OS", "Version", "Notes"],
        [["Windows 10 Pro", "21H2 or later", "For development and small deployments"],
         ["Windows 11 Pro", "Any", "Supported"],
         ["Windows Server", "2019 or 2022", "Recommended for production deployment"]],
        widths=[4, 4, 9])

    doc.add_heading("Software Prerequisites", level=3)
    tbl(doc, ["Software", "Version", "Purpose"],
        [[".NET Runtime", "8.0 (LTS)", "Run ASP.NET Core Web API"],
         ["ASP.NET Core Runtime", "8.0", "HTTP request pipeline"],
         ["SQL Server", "2019 or 2022 (Express/Standard/Enterprise)", "Database engine"],
         ["SQL Server Management Studio", "19.x (optional)", "DB administration"],
         ["Node.js", "18.x LTS or 20.x LTS", "Build React dashboard (dev only)"],
         ["npm", "9.x+ (bundled with Node.js)", "Package management (dev only)"],
         ["Git", "2.x", "Source control (dev only)"]],
        widths=[5, 5, 7])

    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("Note: "); r.font.bold = True
    p.add_run(
        "For production deployment, the React dashboard is built into static files (HTML/CSS/JS) and "
        "served by ASP.NET Core or IIS. Node.js is only required during development or build time."
    )

    # 3.2 Agent PCs
    doc.add_heading("3.2 Factory Agent PC Requirements", level=2)
    doc.add_paragraph("Each factory Machine Controller (MC) runs the Agent application.")

    doc.add_heading("Hardware — Minimum", level=3)
    tbl(doc, ["Component", "Minimum", "Notes"],
        [["Processor", "Any x86/x64 CPU", "Agent is lightweight (~5 MB binary)"],
         ["RAM", "2 GB", "Agent uses < 50 MB"],
         ["Storage", "500 MB free", "For agent binary, models, config, logs"],
         ["Network", "100 Mbps Ethernet", "Must reach server IP"]],
        widths=[3, 5, 9])

    doc.add_heading("Operating System", level=3)
    tbl(doc, ["OS", "Version", "Notes"],
        [["Windows 10", "Any (64-bit preferred)", "Required — Agent uses Win32 API"],
         ["Windows 11", "Any", "Supported"],
         ["Windows 7/8", "SP1", "May work but untested"]],
        widths=[4, 5, 8])

    doc.add_heading("Software Prerequisites", level=3)
    tbl(doc, ["Software", "Version", "Purpose"],
        [["Visual C++ Redistributable", "2022 (v143)", "Runtime for C++ Agent"],
         ["Windows SDK", "10.0 (included in Windows 10+)", "WinHTTP, Winsock libraries"]],
        widths=[5, 5, 7])

    # 3.3 Client
    doc.add_heading("3.3 Client (Dashboard) PC Requirements", level=2)
    doc.add_paragraph("Any PC or device with a modern web browser can access the dashboard.")

    tbl(doc, ["Component", "Requirement"],
        [["Browser", "Google Chrome 90+, Microsoft Edge 90+, or Mozilla Firefox 90+"],
         ["Network", "LAN connectivity to the server PC"],
         ["Display", "1280×720 minimum resolution (1920×1080 recommended)"],
         ["JavaScript", "Must be enabled"],
         ["RAM", "4 GB minimum (for smooth Plotly.js chart rendering)"]],
        widths=[3, 14])

    # 3.4 Network
    doc.add_heading("3.4 Network Requirements", level=2)
    tbl(doc, ["Requirement", "Details"],
        [["Topology", "All components on the same factory LAN (or routable subnets)"],
         ["Bandwidth", "1 Gbps recommended (model files can be up to 500 MB)"],
         ["Ports Required", "Server API: 5000 (HTTP) or 7001 (HTTPS). Dashboard dev: 5173"],
         ["Firewall", "Server port must be open for inbound HTTP from all Agent PCs and client browsers"],
         ["DNS", "Not required — IP-based addressing supported"],
         ["Internet", "Not required after initial software installation"],
         ["Latency", "< 10 ms LAN latency expected"]],
        widths=[4, 13])

    doc.add_page_break()

    # 3.5 Software Dependencies — Server
    doc.add_heading("3.5 Software Dependencies — Server (Backend)", level=2)
    tbl(doc, ["Package", "Version", "License", "Purpose"],
        [["Microsoft.NET", "8.0", "MIT", "Runtime framework"],
         ["ASP.NET Core", "8.0", "MIT", "Web API framework"],
         ["Entity Framework Core", "8.0.0", "MIT", "ORM / database access"],
         ["EF Core SQL Server Provider", "8.0.0", "MIT", "SQL Server connectivity"],
         ["Newtonsoft.Json", "13.0.3", "MIT", "JSON serialization"],
         ["System.IO.Compression", "4.3.0", "MIT", "ZIP archive handling"]],
        widths=[5, 2.5, 2, 7.5])

    # 3.6 Software Dependencies — Frontend
    doc.add_heading("3.6 Software Dependencies — Frontend (Dashboard)", level=2)
    tbl(doc, ["Package", "Version", "Purpose"],
        [["React", "18.2.0", "UI component library"],
         ["React DOM", "18.2.0", "DOM rendering"],
         ["React Router DOM", "6.20.0", "Client-side routing"],
         ["TypeScript", "5.2.2", "Type safety"],
         ["Vite", "5.0.8", "Build tool and dev server"],
         ["Axios", "1.6.2", "HTTP client for API calls"],
         ["Plotly.js", "3.3.1", "Interactive charts (Gantt, bar)"],
         ["Framer Motion", "12.23.26", "Animations and transitions"],
         ["Lucide React", "0.294.0", "Icon library"],
         ["PrismJS", "1.30.0", "Syntax highlighting in editor"],
         ["pako", "2.1.0", "GZIP compression (browser)"],
         ["Zod", "4.3.6", "Runtime schema validation"],
         ["@tanstack/react-virtual", "3.13.18", "Virtualized lists"],
         ["react-simple-code-editor", "0.14.1", "Code editor component"],
         ["clsx", "2.0.0", "Conditional CSS class utility"]],
        widths=[5, 3, 9])

    # 3.7 Software Dependencies — Agent
    doc.add_heading("3.7 Software Dependencies — Agent (C++)", level=2)
    tbl(doc, ["Component", "Version / Details", "Purpose"],
        [["MSVC Toolset", "v143 (Visual Studio 2022) / v145", "C++ compiler"],
         ["C++ Standard", "C++20", "Language standard"],
         ["Windows SDK", "10.0", "Win32 API, WinHTTP, Winsock"],
         ["nlohmann/json", "Header-only (bundled)", "JSON parsing"],
         ["zlib", "1.3.1 (bundled source)", "GZIP compression for images"],
         ["STB Image", "Header-only (bundled)", "Image processing"],
         ["Win32 Subsystem", "Windows", "System tray application"],
         ["Linked Libraries", "winhttp.lib, comctl32.lib, ws2_32.lib", "HTTP, UI controls, sockets"]],
        widths=[4, 5, 8])

    # 3.8 Database
    doc.add_heading("3.8 Database Requirements", level=2)
    tbl(doc, ["Requirement", "Details"],
        [["Engine", "Microsoft SQL Server 2019 or 2022"],
         ["Edition", "Express (free, for up to 10 GB). Standard/Enterprise for production"],
         ["Database Name", "FactoryMonitoringDB"],
         ["Max Connection Pool", "300 (configured in appsettings.json)"],
         ["MARS", "Enabled (MultipleActiveResultSets=True)"],
         ["Authentication", "Windows Authentication (Trusted_Connection)"],
         ["Estimated DB Size", "5-20 GB depending on number of models and version history"],
         ["Tables", "9 tables (FactoryMCs, ConfigFiles, Models, ModelFiles, ModelVersions, LineTargetModels, ModelDistributions, AgentCommands, SystemLogs)"],
         ["Backup", "Regular SQL Server backup recommended"]],
        widths=[4, 13])

    doc.add_page_break()

    # ═══════════════════════════════════════════════════
    # 4. SYSTEM ARCHITECTURE
    # ═══════════════════════════════════════════════════
    doc.add_heading("4. System Architecture", level=1)

    doc.add_heading("4.1 Architecture Diagram", level=2)
    doc.add_paragraph(
        "The system follows a client-server architecture with four main components communicating "
        "over HTTP REST on a factory LAN."
    )
    # Text-based diagram
    arch_text = (
        "┌──────────────────────────────────────────────────────────────┐\n"
        "│                    FACTORY NETWORK (LAN)                     │\n"
        "│                                                              │\n"
        "│  ┌──────────┐  ┌──────────┐  ┌──────────┐                   │\n"
        "│  │ Agent PC  │  │ Agent PC  │  │ Agent PC  │  ... (N MCs)    │\n"
        "│  │  (C++)    │  │  (C++)    │  │  (C++)    │                  │\n"
        "│  └────┬──────┘  └────┬──────┘  └────┬──────┘                 │\n"
        "│       └──────────────┼──────────────┘                        │\n"
        "│                      │ HTTP REST (Port 5000)                 │\n"
        "│                      ▼                                       │\n"
        "│          ┌───────────────────────┐                           │\n"
        "│          │   ASP.NET Core API    │←── Browser (Dashboard)    │\n"
        "│          │   (.NET 8.0)          │    React + TypeScript     │\n"
        "│          └───────────┬───────────┘    (Port 5173 dev)       │\n"
        "│                      │                                       │\n"
        "│                      ▼                                       │\n"
        "│          ┌───────────────────────┐                           │\n"
        "│          │  SQL Server Database  │                           │\n"
        "│          │  (9 tables)           │                           │\n"
        "│          └───────────────────────┘                           │\n"
        "└──────────────────────────────────────────────────────────────┘"
    )
    p = doc.add_paragraph()
    r = p.add_run(arch_text)
    r.font.name = "Consolas"; r.font.size = Pt(8)

    doc.add_heading("4.2 Component Overview", level=2)
    tbl(doc, ["Component", "Technology", "Deployment", "Port"],
        [["Factory Agent", "C++20 / Win32 / WinHTTP", "Each factory MC (system tray app)", "N/A (client)"],
         ["Web API", "ASP.NET Core 8.0 / EF Core 8.0", "Central server PC", "5000 (HTTP)"],
         ["Dashboard", "React 18 / TypeScript / Vite", "Browser (served by API or dev server)", "5173 (dev)"],
         ["Database", "SQL Server 2019+", "Central server PC (same or separate)", "1433 (default)"]],
        widths=[3, 5, 5.5, 3.5])

    doc.add_heading("4.3 Communication Protocols", level=2)
    tbl(doc, ["Path", "Protocol", "Format", "Description"],
        [["Agent → Server", "HTTP REST", "JSON", "Registration, heartbeat, sync, upload"],
         ["Dashboard → Server", "HTTP REST", "JSON + Multipart", "CRUD, file upload/download"],
         ["Server → Agent", "Pull (via heartbeat)", "JSON", "Agent polls for pending commands"],
         ["Server → Database", "TCP (EF Core)", "SQL / TDS", "Data persistence"]],
        widths=[3.5, 3, 3, 7.5])

    doc.add_heading("4.4 Data Flow", level=2)
    flows = [
        ("Agent Registration", "Agent → POST /api/agent/register → Server creates FactoryMCs record in DB."),
        ("Heartbeat", "Agent → POST /api/agent/heartbeat → Server updates IsOnline, LastHeartbeat. Server returns pending commands."),
        ("Model Distribution", "Dashboard → POST /api/modellibrary/apply → Server creates ModelDistribution + AgentCommand records → Agent downloads model on next heartbeat."),
        ("Log Analysis", "Dashboard → POST /api/loganalyzer/analyze → Server sends command to Agent → Agent uploads log → Server parses and returns Gantt data."),
        ("Config Update", "Dashboard → POST /api/mc/{id}/config/update → Server stores pending update → Agent picks up on next heartbeat → Agent applies config."),
    ]
    for name, desc in flows:
        p = doc.add_paragraph()
        r = p.add_run(f"{name}: "); r.font.bold = True
        p.add_run(desc)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════
    # 5. FUNCTIONAL REQUIREMENTS (condensed)
    # ═══════════════════════════════════════════════════
    doc.add_heading("5. Functional Requirements", level=1)

    fr_sections = [
        ("5.1 Real-Time Monitoring", [
            ["FR-MON-001", "Display all PCs grouped by production line with online/offline indicators", "High"],
            ["FR-MON-002", "Filter PCs by model version; support Grid and List view modes", "High"],
            ["FR-MON-003", "Show model compliance per line (how many PCs match target model)", "High"],
            ["FR-MON-004", "Aggregate statistics: total PCs, online, offline, running apps", "Medium"],
            ["FR-MON-005", "Provide 'Manage Models' button per line for model distribution modal", "High"],
        ]),
        ("5.2 PC Management", [
            ["FR-PC-001", "Detailed PC view with all metadata (line, MC#, IP, version, paths, dates)", "High"],
            ["FR-PC-002", "Upload, download, change, and delete models on individual PCs", "High"],
            ["FR-PC-003", "Upload and download configuration files for PCs", "High"],
            ["FR-PC-004", "Edit PC metadata and delete PC registrations", "Medium"],
        ]),
        ("5.3 Model Library & Distribution", [
            ["FR-LIB-001", "Centralized model repository with upload, download, search, delete", "High"],
            ["FR-LIB-002", "Distribute models to: all PCs, specific version, specific line, or selected PCs", "High"],
            ["FR-LIB-003", "Set and track Line Target Models per line per version", "High"],
            ["FR-LIB-004", "Handle offline PCs during distribution with alert and proceed option", "Medium"],
        ]),
        ("5.4 Model Editor & Version History", [
            ["FR-EDIT-001", "In-browser editor with file tree, tabs, syntax highlighting, undo/redo", "High"],
            ["FR-EDIT-002", "Side-by-side diff view with line-level and word-level highlighting", "High"],
            ["FR-EDIT-003", "Auto-create version snapshots on save; display version timeline", "High"],
            ["FR-EDIT-004", "Revert to any previous version; XML visual editor for parameter editing", "High"],
        ]),
        ("5.5 Log Analyser", [
            ["FR-LOG-001", "Browse PC log folder structure and fetch files on demand (cached)", "High"],
            ["FR-LOG-002", "Parse logs into interactive Gantt charts and barrel execution charts", "High"],
            ["FR-LOG-003", "Display inspection images for NG operations with lazy-loading", "Medium"],
            ["FR-LOG-004", "View raw log content and download log files", "Medium"],
        ]),
        ("5.6 Agent Communication", [
            ["FR-AGT-001", "Auto-registration with first-run dialog; periodic heartbeat", "High"],
            ["FR-AGT-002", "Sync model inventory, download pending distributions", "High"],
            ["FR-AGT-003", "Upload/download configs; apply pending config updates", "High"],
            ["FR-AGT-004", "Sync log structure; upload logs and inspection images on demand", "Medium"],
            ["FR-AGT-005", "Run as Windows system tray application", "High"],
        ]),
    ]

    for title, rows in fr_sections:
        doc.add_heading(title, level=2)
        tbl(doc, ["ID", "Requirement", "Priority"], rows, widths=[3, 11, 3])
        doc.add_paragraph()

    doc.add_page_break()

    # ═══════════════════════════════════════════════════
    # 6. NON-FUNCTIONAL REQUIREMENTS
    # ═══════════════════════════════════════════════════
    doc.add_heading("6. Non-Functional Requirements", level=1)
    tbl(doc, ["ID", "Category", "Requirement"],
        [["NFR-001", "Performance", "Dashboard loads within 3 seconds on LAN"],
         ["NFR-002", "Performance", "Model upload supports files up to 500 MB (configurable)"],
         ["NFR-003", "Performance", "Log cache uses LRU eviction, limited to 100 MB (configurable)"],
         ["NFR-004", "Performance", "API timeout: 10 seconds default; log operations: 30-180 sec adaptive"],
         ["NFR-005", "Reliability", "PCs auto-marked offline after 1 minute without heartbeat (configurable)"],
         ["NFR-006", "Reliability", "Concurrent log requests deduplicated to prevent duplicate agent calls"],
         ["NFR-007", "Usability", "Consistent color-coding: green = healthy, red/amber = issues"],
         ["NFR-008", "Usability", "Light and Dark theme support"],
         ["NFR-009", "Usability", "Responsive UI for modern browsers (Chrome 90+, Edge 90+, Firefox 90+)"],
         ["NFR-010", "Maintainability", "Backend: Controllers / Services / Models separation"],
         ["NFR-011", "Maintainability", "Frontend: Pages / Components / Services / Types separation"],
         ["NFR-012", "Scalability", "Database indexes on frequently queried columns"],
         ["NFR-013", "Scalability", "SQL connection pool: 300 max connections"],
         ["NFR-014", "Security", "File path inputs validated against path traversal attacks"],
         ["NFR-015", "Security", "Trusted network model (no auth — factory LAN only)"]],
        widths=[2.5, 3, 11.5])

    doc.add_page_break()

    # ═══════════════════════════════════════════════════
    # 7. DATABASE SCHEMA SUMMARY
    # ═══════════════════════════════════════════════════
    doc.add_heading("7. Database Schema Summary", level=1)
    doc.add_paragraph("Database: FactoryMonitoringDB — SQL Server. 9 tables total.")
    tbl(doc, ["Table", "Purpose", "Key Columns", "Relationships"],
        [["FactoryMCs", "PC registrations", "MCId, LineNumber, MCNumber, IPAddress, ModelVersion, IsOnline", "1→N Models, 1→1 ConfigFiles, 1→N AgentCommands"],
         ["ConfigFiles", "PC configuration content", "ConfigId, MCId, ConfigContent, PendingUpdate", "FK → FactoryMCs (CASCADE)"],
         ["Models", "Models discovered on PCs", "ModelId, MCId, ModelName, IsCurrentModel", "FK → FactoryMCs (CASCADE)"],
         ["ModelFiles", "Centralized model repository", "ModelFileId, ModelName, FileData (ZIP binary), FileSize", "1→N ModelDistributions, 1→N ModelVersions"],
         ["ModelVersions", "Version history snapshots", "ModelVersionId, ModelFileId, VersionNumber, FileData", "FK → ModelFiles"],
         ["LineTargetModels", "Target model per line+version", "LineTargetModelId, LineNumber, ModelVersion, TargetModelName", "Standalone (no FK)"],
         ["ModelDistributions", "Deployment tracking", "DistributionId, ModelFileId, MCId, Status", "FK → ModelFiles, FK → FactoryMCs"],
         ["AgentCommands", "Remote command queue", "CommandId, MCId, CommandType, Status", "FK → FactoryMCs (CASCADE)"],
         ["SystemLogs", "Audit trail", "LogId, MCId, Action, ActionType, Timestamp", "FK → FactoryMCs (SET NULL)"]],
        widths=[3, 3, 6, 5])

    doc.add_page_break()

    # ═══════════════════════════════════════════════════
    # 8. API ENDPOINT SUMMARY
    # ═══════════════════════════════════════════════════
    doc.add_heading("8. API Endpoint Summary", level=1)
    doc.add_paragraph("Base URL: http://<server-ip>:5000/api")
    doc.add_paragraph("Total: 50+ endpoints across 14 controllers.")
    doc.add_paragraph()

    api_groups = [
        ("Dashboard & PC Management (7 endpoints)", [
            ["GET", "/api/versions", "List model versions"],
            ["GET", "/api/lines", "List line numbers"],
            ["GET", "/api/pcs", "Get PCs (filter by version/line)"],
            ["GET", "/api/pcs/{id}", "Get single PC details"],
            ["GET", "/api/stats", "Aggregate statistics"],
            ["PUT", "/api/mc/{id}", "Update PC metadata"],
            ["DELETE", "/api/mc/{id}", "Delete PC"],
        ]),
        ("Model Library (10 endpoints)", [
            ["GET", "/api/modellibrary", "List all library models"],
            ["POST", "/api/modellibrary/upload", "Upload model (multipart)"],
            ["GET", "/api/modellibrary/{id}/download", "Download model"],
            ["DELETE", "/api/modellibrary/{id}", "Delete model"],
            ["POST", "/api/modellibrary/apply", "Distribute model"],
            ["GET", "/api/modellibrary/{id}/structure", "Get file tree"],
            ["GET", "/api/modellibrary/{id}/file", "Get file content"],
            ["POST", "/api/modellibrary/{id}/files/save", "Bulk save files"],
            ["GET", "/api/modellibrary/{id}/history", "Change history"],
            ["POST", "/api/modellibrary/{id}/versions/{v}/revert", "Revert version"],
        ]),
        ("Agent Communication (12 endpoints)", [
            ["POST", "/api/agent/register", "Register agent"],
            ["POST", "/api/agent/heartbeat", "Heartbeat + poll commands"],
            ["POST", "/api/agent/syncmodels", "Sync model list"],
            ["POST", "/api/agent/updateconfig", "Upload config"],
            ["GET", "/api/agent/getconfigupdate", "Get pending config"],
            ["POST", "/api/agent/commandresult", "Report command result"],
            ["POST", "/api/agent/uploadmodel", "Upload model to server"],
            ["GET", "/api/agent/downloadmodel", "Download model"],
            ["POST", "/api/agent/synclogs", "Sync log structure"],
            ["POST", "/api/agent/uploadlog", "Upload log content"],
            ["POST", "/api/agent/uploadimages", "Upload images (base64)"],
            ["POST", "/api/agent/uploadimagesbinary", "Upload images (binary)"],
        ]),
        ("Log Analyser (7 endpoints)", [
            ["GET", "/api/loganalyzer/structure/{mcId}", "Log folder tree"],
            ["POST", "/api/loganalyzer/content", "Get log content"],
            ["POST", "/api/loganalyzer/analyze", "Parse log file"],
            ["POST", "/api/loganalyzer/download", "Download log"],
            ["GET", "/api/loganalyzer/images/{mcId}", "Inspection images"],
            ["GET", "/api/loganalyzer/image/{id}", "Raw image"],
            ["GET", "/api/loganalyzer/fetch-image/{mcId}", "Lazy-load from agent"],
        ]),
    ]

    for title, rows in api_groups:
        doc.add_heading(title, level=2)
        tbl(doc, ["Method", "Endpoint", "Description"], rows, widths=[2, 7.5, 7.5])
        doc.add_paragraph()

    doc.add_page_break()

    # ═══════════════════════════════════════════════════
    # 9. DEPLOYMENT & INSTALLATION
    # ═══════════════════════════════════════════════════
    doc.add_heading("9. Deployment & Installation Guide", level=1)

    doc.add_heading("9.1 Database Setup", level=2)
    steps = [
        "1. Install SQL Server 2019/2022 (Express edition is sufficient for small deployments).",
        "2. Open SQL Server Management Studio (SSMS) and connect to the instance.",
        "3. Execute database scripts in order:",
        "   • db/01_CreateDatabase.sql — Creates FactoryMonitoringDB",
        "   • db/02_CreateTables.sql — Creates all 9 tables with constraints",
        "   • db/03_SeedData.sql — (Optional) Insert sample data",
        "4. Verify tables are created: FactoryMCs, ConfigFiles, Models, ModelFiles, ModelVersions, LineTargetModels, ModelDistributions, AgentCommands, SystemLogs.",
    ]
    for s in steps:
        doc.add_paragraph(s)

    doc.add_heading("9.2 Backend Server Setup", level=2)
    steps = [
        "1. Install .NET 8.0 Runtime (or SDK for development).",
        "2. Clone the repository or copy FactoryMonitoringWeb/ folder to server.",
        "3. Edit appsettings.json — update ConnectionStrings.DefaultConnection with your SQL Server instance name.",
        '4. Run: dotnet run --project FactoryMonitoringWeb/ --urls "http://0.0.0.0:5000"',
        "5. Verify: Open http://<server-ip>:5000/api/stats in a browser — should return JSON.",
        "6. Ensure Windows Firewall allows inbound TCP on port 5000.",
    ]
    for s in steps:
        doc.add_paragraph(s)

    doc.add_heading("9.3 Dashboard Setup (Development)", level=2)
    steps = [
        "1. Install Node.js 18.x or 20.x LTS.",
        "2. Navigate to factory-react-ui/ directory.",
        "3. Run: npm install",
        "4. Edit src/services/api.ts — set baseURL to http://<server-ip>:5000/api",
        "5. Run: npm run dev -- --host",
        "6. Access dashboard at http://<server-ip>:5173",
    ]
    for s in steps:
        doc.add_paragraph(s)

    doc.add_heading("9.4 Dashboard Setup (Production)", level=2)
    steps = [
        "1. Run: npm run build (outputs to dist/ folder)",
        "2. Copy dist/ contents to the ASP.NET Core wwwroot/ folder or configure IIS to serve static files.",
        "3. No Node.js required on production server.",
    ]
    for s in steps:
        doc.add_paragraph(s)

    doc.add_heading("9.5 Agent Deployment", level=2)
    steps = [
        "1. Build FactoryAgent.sln using Visual Studio 2022 (Release x64 configuration).",
        "2. Copy the output binary (FactoryAgent.exe) to each factory PC.",
        "3. Ensure Visual C++ 2022 Redistributable (x64) is installed on the PC.",
        "4. Run FactoryAgent.exe — a registration dialog will appear on first run.",
        "5. Enter: Server URL (http://<server-ip>:5000), Line Number, MC Number.",
        "6. Agent will register with the server and begin heartbeat loop.",
        "7. Optionally, add FactoryAgent.exe to Windows Startup for auto-start.",
    ]
    for s in steps:
        doc.add_paragraph(s)

    doc.add_heading("9.6 Configuration Parameters", level=2)
    tbl(doc, ["Parameter", "File", "Default", "Description"],
        [["ConnectionString", "appsettings.json", "(local)\\SQLEXPRESS", "SQL Server connection string"],
         ["HeartbeatTimeoutMinutes", "appsettings.json", "1", "Minutes before PC marked offline"],
         ["MaxUploadSizeMB", "appsettings.json", "500", "Max file upload size"],
         ["BaseTimeoutSeconds", "appsettings.json", "30", "Base timeout for log requests"],
         ["TimeoutPerMBSeconds", "appsettings.json", "5", "Additional timeout per MB of log"],
         ["MaxTimeoutSeconds", "appsettings.json", "180", "Maximum log request timeout"],
         ["CacheSizeLimitMB", "appsettings.json", "100", "LRU log cache size limit"],
         ["API Base URL", "api.ts", "http://localhost:5000", "Dashboard → Server URL"]],
        widths=[4, 3, 3.5, 6.5])

    doc.add_page_break()

    # ═══════════════════════════════════════════════════
    # 10. GLOSSARY
    # ═══════════════════════════════════════════════════
    doc.add_heading("10. Glossary", level=1)
    tbl(doc, ["Term", "Definition"],
        [["MC", "Machine Controller — A factory PC on a production line"],
         ["Line", "Production line containing multiple MCs"],
         ["Model Version", "Software/hardware version string (e.g., '3.5', '4.0')"],
         ["Model", "Machine learning model package stored as a ZIP archive"],
         ["Target Model", "Designated standard model for a production line"],
         ["Compliance", "State where PC's active model matches line's target"],
         ["Agent", "C++ desktop application running on each factory PC"],
         ["Heartbeat", "Periodic HTTP call from agent to server"],
         ["NG", "No Good — Failed inspection result"],
         ["Barrel", "A unit of work/inspection in factory process"],
         ["Distribution", "Deploying a model from library to one or more PCs"],
         ["LAN", "Local Area Network"],
         ["SRS", "Software Requirements Specification"],
         ["WBS", "Work Breakdown Structure"],
         ["EF Core", "Entity Framework Core — .NET ORM"],
         ["ORM", "Object-Relational Mapping"]],
        widths=[4, 13])

    # ── Footer ──
    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("— End of Document —"); r.font.italic = True
    r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Factory Monitoring System — SRS v1.0 — February 2026")
    r.font.size = Pt(10); r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Contributors: Divyansh, Bhavik, Praganya")
    r.font.size = Pt(10); r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.save(OUT)
    print(f"Done → {OUT}")

if __name__ == "__main__":
    build()
