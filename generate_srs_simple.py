"""
Simple SRS Document — Hardware Requirements + Feature SRS Table
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

OUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "SRS_FactoryMonitoring_Final.docx")

def shade(cell, color):
    cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>'))

def tbl(doc, hdrs, rows, widths=None, hdr_bg="1F4E79"):
    t = doc.add_table(rows=1+len(rows), cols=len(hdrs))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(hdrs):
        t.rows[0].cells[j].text = h
    for i, rd in enumerate(rows):
        for j, v in enumerate(rd):
            t.rows[i+1].cells[j].text = v
    # Style header
    for c in t.rows[0].cells:
        shade(c, hdr_bg)
        for p in c.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor.from_string("FFFFFF")
                r.font.bold = True; r.font.size = Pt(10)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.space_before = Pt(3)
    # Style data rows
    for i, row in enumerate(t.rows[1:], 1):
        bg = "F2F7FB" if i % 2 == 0 else "FFFFFF"
        for c in row.cells:
            shade(c, bg)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.space_before = Pt(3)
    if widths:
        for row in t.rows:
            for j, w in enumerate(widths):
                row.cells[j].width = Cm(w)
    return t

def build():
    doc = Document()
    s = doc.sections[0]
    s.page_width = Inches(8.5); s.page_height = Inches(11)
    s.top_margin = Cm(2); s.bottom_margin = Cm(2)
    s.left_margin = Cm(2); s.right_margin = Cm(2)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'; style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    for lv in range(1, 4):
        hs = doc.styles[f'Heading {lv}']
        hs.font.name = 'Calibri'; hs.font.bold = True
        hs.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        hs.font.size = Pt([0, 20, 15, 12][lv])

    # ── COVER PAGE ──
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Software Requirements Specification"); r.font.size = Pt(26); r.font.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Factory Monitoring System"); r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

    doc.add_paragraph()
    for lbl, val in [("Date:", "February 13, 2026"), ("Version:", "1.0"),
                      ("Team:", "Divyansh, Bhavik, Praganya")]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rl = p.add_run(lbl + " "); rl.font.bold = True; rl.font.size = Pt(11)
        p.add_run(val).font.size = Pt(11)

    doc.add_page_break()

    # ══════════════════════════════════════════
    # SECTION 1: PROJECT OVERVIEW
    # ══════════════════════════════════════════
    doc.add_heading("1. Project Overview", level=1)
    doc.add_paragraph(
        "The Factory Monitoring System is a centralized solution for remotely monitoring factory "
        "production-line PCs (Machine Controllers), distributing inspection models, managing "
        "configurations, and analyzing operation logs — all from a web dashboard."
    )
    doc.add_paragraph("The system has 4 main components:")
    for item in [
        "Factory Agent — A C++ application running on each factory PC that communicates with the server.",
        "Web API Server — An ASP.NET Core backend that handles all data and operations.",
        "Web Dashboard — A React-based browser application for engineers to monitor and manage everything.",
        "SQL Server Database — Stores PC records, model files, configurations, commands, and logs.",
    ]:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()

    # ══════════════════════════════════════════
    # SECTION 2: HARDWARE REQUIREMENTS
    # ══════════════════════════════════════════
    doc.add_heading("2. Hardware Requirements", level=1)

    doc.add_heading("2.1 Server PC (Hosts Web API + Database)", level=2)
    tbl(doc, ["Component", "Minimum Requirement", "Recommended"],
        [["Processor", "Standard desktop/server processor (x64)", "Quad-core or higher"],
         ["RAM", "8 GB", "16 GB"],
         ["Storage (HDD/SSD)", "50 GB free (HDD acceptable, SSD preferred)", "100 GB+ SSD for faster DB performance"],
         ["Architecture", "64-bit (x64)", "64-bit (x64)"],
         ["Network", "100 Mbps Ethernet (LAN)", "1 Gbps Ethernet"],
         ["OS", "Windows 10 Pro (21H2+)", "Windows Server 2019 / 2022"]],
        widths=[3.5, 6.5, 6.5])

    doc.add_paragraph()
    doc.add_heading("2.2 Factory PC (Runs Agent)", level=2)
    tbl(doc, ["Component", "Minimum Requirement", "Notes"],
        [["Processor", "Standard desktop processor (x86/x64)", "Agent is very lightweight (~5 MB binary)"],
         ["RAM", "2 GB", "Agent uses less than 50 MB memory"],
         ["Storage (HDD/SSD)", "500 MB free (HDD or SSD)", "For agent binary, models, configs, and logs"],
         ["Architecture", "32-bit or 64-bit", "64-bit preferred"],
         ["Network", "100 Mbps Ethernet", "Must be able to reach server IP on LAN"],
         ["OS", "Windows 10 or later", "Required — agent uses Win32 API"]],
        widths=[3.5, 6.5, 6.5])

    doc.add_paragraph()
    doc.add_heading("2.3 Client PC (Accesses Dashboard via Browser)", level=2)
    tbl(doc, ["Component", "Requirement"],
        [["Browser", "Chrome 90+, Edge 90+, or Firefox 90+"],
         ["Network", "LAN access to server PC"],
         ["Display", "1280×720 minimum (1920×1080 recommended)"],
         ["RAM", "4 GB minimum"]],
        widths=[3.5, 13])

    doc.add_paragraph()
    doc.add_heading("2.4 Network", level=2)
    tbl(doc, ["Requirement", "Details"],
        [["Topology", "All PCs on same factory LAN (or routable subnets)"],
         ["Bandwidth", "100 Mbps sufficient (model files are under 10 MB each)"],
         ["Server Ports", "Port 5000 (HTTP API), Port 1433 (SQL Server)"],
         ["Firewall", "Server port 5000 must be open for inbound connections from agent PCs and client browsers"],
         ["Internet", "Required only during initial setup (for installing .NET, SQL Server, Node.js). Not needed during normal operation."]],
        widths=[3.5, 13])

    doc.add_page_break()

    # ══════════════════════════════════════════
    # SECTION 3: SOFTWARE REQUIREMENTS
    # ══════════════════════════════════════════
    doc.add_heading("3. Software Requirements", level=1)

    doc.add_heading("3.1 Server Software", level=2)
    tbl(doc, ["Software", "Version", "Purpose"],
        [[".NET Runtime", "8.0", "Runs the Web API server"],
         ["SQL Server", "2019 or 2022 (Express is fine)", "Database engine"],
         ["Node.js", "18.x or 20.x LTS", "Only needed during development to build the dashboard"],
         ["Git", "2.x", "Source control (development only)"]],
        widths=[5, 5.5, 6.5])

    doc.add_paragraph()
    doc.add_heading("3.2 Factory PC Software", level=2)
    tbl(doc, ["Software", "Version", "Purpose"],
        [["Visual C++ Redistributable", "2022 (x64)", "Required to run the Agent"],
         ["Windows 10+", "Any", "OS requirement for Win32 API"]],
        widths=[5, 5.5, 6.5])

    doc.add_paragraph()
    doc.add_heading("3.3 Client Software", level=2)
    tbl(doc, ["Software", "Version"],
        [["Google Chrome / Microsoft Edge / Firefox", "Version 90 or later"],
         ["JavaScript", "Must be enabled in browser"]],
        widths=[10, 7])

    doc.add_page_break()

    # ══════════════════════════════════════════
    # SECTION 4: SRS — FEATURE REQUIREMENTS
    # ══════════════════════════════════════════
    doc.add_heading("4. Software Requirements Specification", level=1)
    doc.add_paragraph(
        "The following table lists all the software requirements for the Factory Monitoring System. "
        "Each SRS entry describes what the feature is, what it does, and its priority."
    )
    doc.add_paragraph()

    tbl(doc, ["SRS ID", "Feature / Module", "Requirement Description", "Priority"],
        [
            ["SRS-001", "Real-Time Dashboard & Monitoring",
             "• Display all factory PCs grouped by production line with online/offline color-coded indicators.\n"
             "• Filter PCs by model version (e.g., 3.5, 4.0) via sidebar navigation tree.\n"
             "• Show per-line statistics: total PCs, online/offline count, and model compliance.\n"
             "• Support Grid and List view modes with collapsible production lines.",
             "High"],

            ["SRS-002", "PC Details & Administration",
             "• Detailed PC view showing: line number, MC number, IP address, model version, config/log/model paths, registration date, last heartbeat, and status.\n"
             "• Edit PC metadata (line, MC number, IP, paths, version).\n"
             "• Delete PC registrations from the system.",
             "High"],

            ["SRS-003", "PC Model & Config Operations",
             "• Upload, download, change, and delete models on individual PCs via the agent.\n"
             "• Download current configuration file from a PC.\n"
             "• Upload and push new configuration files to a PC remotely.",
             "High"],

            ["SRS-004", "Model Library & Distribution",
             "• Centralized Model Library with upload, download, search, and delete of model files (.zip).\n"
             "• Distribute models to: all PCs, a specific version, a specific line, or selected PCs.\n"
             "• Set and track target model per line per version with compliance monitoring.\n"
             "• Manage Models modal per line showing available models, compliance, and apply/delete options.",
             "High"],

            ["SRS-005", "Model Editor & XML Visual Editor",
             "• In-browser code editor with file tree, multi-tab support, and syntax highlighting (XML, JSON, Python, INI, YAML).\n"
             "• Undo/redo, unsaved change tracking, and save files back to the model ZIP.\n"
             "• Side-by-side diff view with line-level and word-level change highlighting.\n"
             "• XML Visual Editor for parameter editing in a structured tree (Groups → Specs → Values) with inline diff and revert.",
             "High"],

            ["SRS-006", "Version History & Rollback",
             "• Auto-create version snapshot on every model save.\n"
             "• View full version timeline and compare any two versions with side-by-side diff.\n"
             "• View parameter-level changes for XML files.\n"
             "• Revert a model to any previous version.",
             "High"],

            ["SRS-007", "Log Analyser & Visualization",
             "• Browse log folder structure of any PC; fetch log files on demand with server-side caching.\n"
             "• Parse logs into interactive Gantt chart timelines and barrel execution breakdown charts.\n"
             "• Display inspection images for failed (NG) operations with lazy-loading and gallery view.\n"
             "• View raw log content and download log files.",
             "High"],

            ["SRS-008", "Agent Registration & Heartbeat",
             "• Auto-registration on first run via dialog (server URL, line number, MC number).\n"
             "• Periodic heartbeat reporting online status and application running state.\n"
             "• PCs automatically marked offline after missed heartbeats.\n"
             "• Runs as a Windows system tray application with minimal resource usage.",
             "High"],

            ["SRS-009", "Agent Sync Operations",
             "• Sync local model inventory with server and download pending model distributions.\n"
             "• Upload/download model files and configuration files on request.\n"
             "• Sync log folder structure, upload log content and inspection images on demand.",
             "High"],

            ["SRS-010", "UI Framework & Navigation",
             "• Collapsible sidebar with links to Dashboard, Model Library, and Log Analyzer.\n"
             "• Version-and-line navigation tree with per-line online/offline counts.\n"
             "• Light and Dark theme toggle.\n"
             "• Client-side routing with 404 handling and user-friendly error toast notifications.",
             "Medium"],
        ],
        widths=[1.8, 3.5, 10, 1.7])

    doc.add_page_break()

    # ══════════════════════════════════════════
    # SECTION 5: NON-FUNCTIONAL REQUIREMENTS
    # ══════════════════════════════════════════
    doc.add_heading("5. Non-Functional Requirements", level=1)

    tbl(doc, ["NFR ID", "Category", "Requirement"],
        [["NFR-001", "Performance & Reliability",
          "• Dashboard loads within 3 seconds on LAN.\n"
          "• Model uploads up to 10 MB supported.\n"
          "• Log cache limited to 100 MB with automatic eviction.\n"
          "• PCs auto-marked offline after 1 minute without heartbeat.\n"
          "• API timeout: 10 seconds default; log operations: up to 180 seconds."],
         ["NFR-002", "Usability & Security",
          "• Light and Dark theme support with consistent color-coding (green = OK, red = issue).\n"
          "• Responsive UI for Chrome 90+, Edge 90+, Firefox 90+.\n"
          "• File path validation to prevent path traversal attacks.\n"
          "• Operates on isolated factory LAN — no internet exposure during normal operation."]],
        widths=[2, 3, 12])

    # ── Footer ──
    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("— End of Document —"); r.font.italic = True
    r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.save(OUT)
    print(f"Done → {OUT}")

if __name__ == "__main__":
    build()
