"""
Opens existing SRS_FactoryMonitoring_Final 3 (1).docx,
inserts Gantt chart tables AFTER the functional SRS table (Table 7)
and AFTER the non-functional SRS table (Table 8).
No other content is modified. No color coding in Gantt cells.
"""
from docx import Document
from docx.shared import Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml, OxmlElement
from copy import deepcopy
import os

SRC = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                   "SRS_FactoryMonitoring_Final 3 (1).docx")
OUT = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                   "SRS_FactoryMonitoring_Final_WithGantt.docx")

MONTHS = ["Mar'26", "Apr'26", "May'26", "Jun'26"]
WEEKS_PER_MONTH = 4
TOTAL_WEEKS = 16

def shade(cell, color):
    cell._tc.get_or_add_tcPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>'))

def set_cell_text(cell, text, size=7, bold=False, align=None):
    # Clear existing
    for p in cell.paragraphs:
        for r in p.runs:
            r.clear()
    cell.text = text
    for p in cell.paragraphs:
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        if align:
            p.alignment = align
        for r in p.runs:
            r.font.size = Pt(size)
            r.font.bold = bold
            r.font.name = 'Calibri'

def build_gantt_table(doc, title_text, tasks):
    """
    Build a Gantt chart table element and a preceding title paragraph.
    tasks = list of (srs_id, feature, start_week, end_week)
    Weeks: 0-3=Mar, 4-7=Apr, 8-11=May, 12-15=Jun
    Returns (title_paragraph_element, table_element)
    """
    # Create title paragraph
    title_para = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '200')
    spacing.set(qn('w:after'), '100')
    pPr.append(spacing)
    title_para.append(pPr)
    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    b = OxmlElement('w:b')
    rPr.append(b)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '22')  # 11pt
    rPr.append(sz)
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '1F4E79')
    rPr.append(color)
    run.append(rPr)
    t = OxmlElement('w:t')
    t.text = title_text
    run.append(t)
    title_para.append(run)

    # Build table using Document API then extract element
    num_cols = 3 + TOTAL_WEEKS  # SRS ID | Feature | Timeline Label | 16 week cols
    num_data_rows = len(tasks)
    num_rows = 2 + num_data_rows  # 2 header rows + data

    tbl = doc.add_table(rows=num_rows, cols=num_cols)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False

    # ── Row 0: Main headers ──
    row0 = tbl.rows[0]
    set_cell_text(row0.cells[0], "SRS ID", size=7, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(row0.cells[1], "Feature / Module", size=7, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(row0.cells[2], "Sub Tasks", size=7, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    shade(row0.cells[0], "1F4E79")
    shade(row0.cells[1], "1F4E79")
    shade(row0.cells[2], "1F4E79")
    for p in row0.cells[0].paragraphs:
        for r in p.runs:
            r.font.color.rgb = parse_xml(f'<w:color {nsdecls("w")} w:val="FFFFFF"/>') or None
    # Set white font for header text cells
    for ci in range(3):
        for p in row0.cells[ci].paragraphs:
            for r in p.runs:
                rPr = r._element.get_or_add_rPr()
                c_elem = OxmlElement('w:color')
                c_elem.set(qn('w:val'), 'FFFFFF')
                rPr.append(c_elem)

    # Merge month cells
    for mi, month in enumerate(MONTHS):
        start_col = 3 + mi * WEEKS_PER_MONTH
        end_col = start_col + WEEKS_PER_MONTH - 1
        cell_start = row0.cells[start_col]
        cell_end = row0.cells[end_col]
        cell_start.merge(cell_end)
        set_cell_text(cell_start, month, size=7, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade(cell_start, "1F4E79")
        for p in cell_start.paragraphs:
            for r in p.runs:
                rPr = r._element.get_or_add_rPr()
                c_elem = OxmlElement('w:color')
                c_elem.set(qn('w:val'), 'FFFFFF')
                rPr.append(c_elem)

    # ── Row 1: Week sub-headers ──
    row1 = tbl.rows[1]
    set_cell_text(row1.cells[0], "", size=7)
    set_cell_text(row1.cells[1], "", size=7)
    set_cell_text(row1.cells[2], "", size=7)
    shade(row1.cells[0], "D6E4F0")
    shade(row1.cells[1], "D6E4F0")
    shade(row1.cells[2], "D6E4F0")
    for wi in range(TOTAL_WEEKS):
        cell = row1.cells[3 + wi]
        set_cell_text(cell, f"W{(wi % 4) + 1}", size=6, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade(cell, "D6E4F0")

    # ── Data rows ──
    for ti, (srs_id, feature, sub_task, sw, ew) in enumerate(tasks):
        row = tbl.rows[2 + ti]
        set_cell_text(row.cells[0], srs_id, size=7)
        set_cell_text(row.cells[1], feature, size=7)
        set_cell_text(row.cells[2], sub_task, size=7)

        # Alternate row bg for text cells
        bg = "F2F7FB" if ti % 2 == 0 else "FFFFFF"
        shade(row.cells[0], bg)
        shade(row.cells[1], bg)
        shade(row.cells[2], bg)

        # Week cells - NO color, just "X" marker for scheduled weeks
        for wi in range(TOTAL_WEEKS):
            cell = row.cells[3 + wi]
            if sw <= wi <= ew:
                set_cell_text(cell, "X", size=6, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            else:
                set_cell_text(cell, "", size=6)
            shade(cell, "FFFFFF")  # Always white bg - no color coding

    # Set column widths
    for row in tbl.rows:
        row.cells[0].width = Cm(1.8)
        row.cells[1].width = Cm(3.5)
        row.cells[2].width = Cm(4.0)
        for wi in range(TOTAL_WEEKS):
            row.cells[3 + wi].width = Cm(0.65)

    # Extract the table element
    tbl_element = tbl._element
    return title_para, tbl_element


def insert_after(parent, ref_element, new_element):
    """Insert new_element right after ref_element in parent."""
    ref_element.addnext(new_element)


def main():
    doc = Document(SRC)
    body = doc.element.body

    # Find Table 7 (functional SRS) and Table 8 (NFR) by index
    all_tables = doc.tables
    srs_table = all_tables[7]  # Table index 7 = functional SRS
    nfr_table = all_tables[8]  # Table index 8 = NFR

    srs_tbl_elem = srs_table._element
    nfr_tbl_elem = nfr_table._element

    # ── Functional SRS Gantt Chart ──
    # 10 SRS items distributed across Mar-Jun 2026
    # Weeks: 0-3=Mar, 4-7=Apr, 8-11=May, 12-15=Jun
    fr_tasks = [
        ("SRS-001", "Real-Time Dashboard", "Dashboard & monitoring setup", 0, 3),
        ("SRS-002", "PC Details & Admin", "PC detail views & admin", 2, 5),
        ("SRS-003", "PC Model & Config Ops", "Model/config operations", 3, 6),
        ("SRS-004", "Model Library & Dist.", "Library, distribution pipeline", 4, 8),
        ("SRS-005", "Model Editor & XML", "Editor, XML visual editor", 5, 9),
        ("SRS-006", "Version History", "Versioning & rollback", 7, 10),
        ("SRS-007", "Log Analyser", "Log parsing & visualization", 3, 8),
        ("SRS-008", "Agent Registration", "Agent registration & heartbeat", 0, 3),
        ("SRS-009", "Agent Sync Ops", "Sync models, configs, logs", 2, 6),
        ("SRS-010", "UI Framework & Nav", "Sidebar, themes, routing", 0, 4),
    ]

    title1, gantt1 = build_gantt_table(doc,
        "Functional Requirements — Work Distribution (March 2026 – June 2026)",
        fr_tasks)

    # ── Non-Functional SRS Gantt Chart ──
    nfr_tasks = [
        ("NFR-001", "Performance & Reliability", "Load time, caching, timeouts", 0, 15),
        ("NFR-002", "Usability & Security", "Themes, color-coding, validation", 0, 15),
    ]

    title2, gantt2 = build_gantt_table(doc,
        "Non-Functional Requirements — Work Distribution (March 2026 – June 2026)",
        nfr_tasks)

    # Now we need to move these tables to the right positions.
    # The build_gantt_table added them at the end of the doc.
    # We need to remove them from end and insert after the SRS/NFR tables.

    # Remove the 4 elements we just added (2 titles + 2 tables) from body end
    # Actually, build_gantt_table only added the table via doc.add_table.
    # The title paragraph was created as raw XML and never added.
    # So we have 2 extra tables at the end.

    # Get the gantt table elements (they are the last 2 tables added)
    all_body_children = list(body)

    # The gantt tables are the last 2 table elements added
    # Remove them from end of body
    body.remove(gantt1)
    body.remove(gantt2)

    # Insert after SRS table: title1 then gantt1
    # Also add a small spacing paragraph after the gantt
    spacer1 = OxmlElement('w:p')
    sp1 = OxmlElement('w:pPr')
    sp1_spacing = OxmlElement('w:spacing')
    sp1_spacing.set(qn('w:after'), '50')
    sp1.append(sp1_spacing)
    spacer1.append(sp1)

    insert_after(body, srs_tbl_elem, spacer1)
    insert_after(body, srs_tbl_elem, gantt1)
    insert_after(body, srs_tbl_elem, title1)

    # Insert after NFR table: title2 then gantt2
    spacer2 = OxmlElement('w:p')
    sp2 = OxmlElement('w:pPr')
    sp2_spacing = OxmlElement('w:spacing')
    sp2_spacing.set(qn('w:after'), '50')
    sp2.append(sp2_spacing)
    spacer2.append(sp2)

    insert_after(body, nfr_tbl_elem, spacer2)
    insert_after(body, nfr_tbl_elem, gantt2)
    insert_after(body, nfr_tbl_elem, title2)

    doc.save(OUT)
    print(f"Done → {OUT}")


if __name__ == "__main__":
    main()
