"""
Generate SRS_FactoryMonitoring_Final.docx with Gantt-chart timeline tables
for Functional & Non-Functional requirements (March 2026 – June 2026).
All other sections remain identical to the original SRS document.
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

OUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "SRS_FactoryMonitoring_Final.docx")

# ── Color constants ──
MONTHS = ["Mar'26", "Apr'26", "May'26", "Jun'26"]
WEEKS_PER_MONTH = 4
TOTAL_WEEKS = len(MONTHS) * WEEKS_PER_MONTH  # 16 weeks

# Colors for Gantt bars (cycle through these)
BAR_COLORS = [
    "4472C4",  # Blue
    "70AD47",  # Green
    "FFC000",  # Yellow/Gold
    "ED7D31",  # Orange
    "5B9BD5",  # Light Blue
    "A5A5A5",  # Grey
    "44546A",  # Dark Blue-Grey
    "BDD7EE",  # Pale Blue
    "C5E0B4",  # Pale Green
    "FFE699",  # Pale Yellow
]

def shade(cell, color):
    cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>'))

def set_cell_font(cell, size=7, bold=False, color_rgb=None):
    for p in cell.paragraphs:
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.space_before = Pt(1)
        for r in p.runs:
            r.font.size = Pt(size)
            r.font.bold = bold
            if color_rgb:
                r.font.color.rgb = color_rgb

def style_tbl(t, hdr_bg="1F4E79", hdr_fg="FFFFFF"):
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for c in t.rows[0].cells:
        shade(c, hdr_bg)
        for p in c.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor.from_string(hdr_fg)
                r.font.bold = True; r.font.size = Pt(9)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
    for i, row in enumerate(t.rows[1:], 1):
        bg = "F2F7FB" if i % 2 == 0 else "FFFFFF"
        for c in row.cells:
            shade(c, bg)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.space_before = Pt(2)

def tbl(doc, hdrs, rows, widths=None, bg="1F4E79"):
    t = doc.add_table(rows=1+len(rows), cols=len(hdrs))
    t.style = 'Table Grid'
    for j, h in enumerate(hdrs):
        t.rows[0].cells[j].text = h
    for i, rd in enumerate(rows):
        for j, v in enumerate(rd):
            t.rows[i+1].cells[j].text = v
    if widths:
        for row in t.rows:
            for j, w in enumerate(widths):
                row.cells[j].width = Cm(w)
    style_tbl(t, hdr_bg=bg)
    return t


def create_gantt_table(doc, title, tasks, color_start=0):
    """
    Create a Gantt chart table like the sample image.
    tasks = list of (functionality, main_task, sub_task, start_week, end_week)
      where start_week and end_week are 0-indexed (0 = Mar W1, 15 = Jun W4)
    """
    doc.add_heading(title, level=2)
    
    # Columns: Functionality | Main Tasks | Sub Tasks | Mar W1-W4 | Apr W1-W4 | May W1-W4 | Jun W1-W4
    num_cols = 3 + TOTAL_WEEKS  # 3 text cols + 16 week cols
    num_rows = 2 + len(tasks)   # 2 header rows + data rows
    
    t = doc.add_table(rows=num_rows, cols=num_cols)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    
    # ── Row 0: Month headers (merged across 4 weeks each) ──
    row0 = t.rows[0]
    # First 3 cells: Functionality, Main Tasks, Sub Tasks
    row0.cells[0].text = "Functionality"
    row0.cells[1].text = "Main Tasks"
    row0.cells[2].text = "Sub Tasks"
    
    # Merge month header cells
    for mi, month in enumerate(MONTHS):
        start_col = 3 + mi * WEEKS_PER_MONTH
        end_col = start_col + WEEKS_PER_MONTH - 1
        # Merge cells for this month
        cell_start = row0.cells[start_col]
        cell_end = row0.cells[end_col]
        cell_start.merge(cell_end)
        cell_start.text = month
    
    # Style row 0
    for c in row0.cells:
        shade(c, "1F4E79")
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.bold = True
                r.font.size = Pt(8)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.space_before = Pt(1)
    
    # ── Row 1: Week headers (W1, W2, W3, W4 repeated) ──
    row1 = t.rows[1]
    row1.cells[0].text = ""
    row1.cells[1].text = ""
    row1.cells[2].text = ""
    for wi in range(TOTAL_WEEKS):
        row1.cells[3 + wi].text = f"W{(wi % 4) + 1}"
    
    for c in row1.cells:
        shade(c, "D6E4F0")
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.size = Pt(7)
                r.font.bold = True
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.space_before = Pt(1)
    
    # ── Data rows ──
    for ti, (func, main, sub, sw, ew) in enumerate(tasks):
        row = t.rows[2 + ti]
        row.cells[0].text = func
        row.cells[1].text = main
        row.cells[2].text = sub
        
        bar_color = BAR_COLORS[(color_start + ti) % len(BAR_COLORS)]
        
        for ci in range(TOTAL_WEEKS):
            cell = row.cells[3 + ci]
            if sw <= ci <= ew:
                shade(cell, bar_color)
            else:
                shade(cell, "FFFFFF")
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(1)
                p.paragraph_format.space_before = Pt(1)
        
        # Style text cells
        for ci in range(3):
            cell = row.cells[ci]
            bg_color = "F2F7FB" if ti % 2 == 0 else "FFFFFF"
            shade(cell, bg_color)
            set_cell_font(cell, size=7)
    
    # ── Set column widths ──
    # Text columns wider, week columns narrow
    text_widths = [Cm(3.0), Cm(3.5), Cm(4.0)]
    week_width = Cm(0.7)
    
    for row in t.rows:
        for ci in range(3):
            row.cells[ci].width = text_widths[ci]
        for ci in range(TOTAL_WEEKS):
            row.cells[3 + ci].width = week_width
    
    doc.add_paragraph()
    return t


def build():
    doc = Document()

    # Page setup - landscape for Gantt charts
    s = doc.sections[0]
    s.page_width = Inches(11); s.page_height = Inches(8.5)  # Landscape
    s.top_margin = Cm(1.5); s.bottom_margin = Cm(1.5)
    s.left_margin = Cm(1.5); s.right_margin = Cm(1.5)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'; style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    for lv in range(1, 4):
        hs = doc.styles[f'Heading {lv}']
        hs.font.name = 'Calibri'
        hs.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        hs.font.bold = True
        hs.font.size = Pt([0, 22, 16, 13][lv])

    # ── COVER PAGE ──
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Software Requirements Specification (SRS)")
    r.font.size = Pt(28); r.font.bold = True; r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Factory Monitoring System")
    r.font.size = Pt(20); r.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Project Timeline & Requirements Schedule")
    r.font.size = Pt(14); r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()
    for label, val in [("Version:", "1.0"), ("Date:", "February 16, 2026"),
                       ("Project Duration:", "December 2025 – June 2026"),
                       ("Prepared By:", "Divyansh, Bhavik, Praganya")]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rl = p.add_run(label + " "); rl.font.bold = True; rl.font.size = Pt(12)
        rl.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        rv = p.add_run(val); rv.font.size = Pt(12); rv.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    doc.add_page_break()

    # ── TABLE OF CONTENTS ──
    doc.add_heading("Table of Contents", level=1)
    toc = [
        "1. Introduction",
        "   1.1 Purpose", "   1.2 Scope", "   1.3 Intended Audience",
        "   1.4 Project Team", "   1.5 Definitions & Abbreviations", "   1.6 References",
        "2. Overall Description",
        "   2.1 Product Perspective", "   2.2 Product Functions", "   2.3 User Classes",
        "   2.4 Constraints & Assumptions",
        "3. System Requirements",
        "   3.1 Server PC Requirements", "   3.2 Factory Agent PC Requirements",
        "   3.3 Client PC Requirements", "   3.4 Network Requirements",
        "   3.5–3.8 Software & Database Dependencies",
        "4. System Architecture",
        "5. Functional Requirements — Project Timeline (Gantt Chart)",
        "6. Non-Functional Requirements — Project Timeline (Gantt Chart)",
        "7. Database Schema Summary",
        "8. API Endpoint Summary",
        "9. Deployment & Installation Guide",
        "10. Glossary",
    ]
    for item in toc:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
        p.runs[0].font.size = Pt(11)

    doc.add_page_break()

    # ═══ 1. INTRODUCTION ═══
    doc.add_heading("1. Introduction", level=1)

    doc.add_heading("1.1 Purpose", level=2)
    doc.add_paragraph(
        "This Software Requirements Specification (SRS) document describes the complete requirements "
        "for the Factory Monitoring System — a full-stack, enterprise-grade solution for remotely "
        "monitoring factory production-line PCs (Machine Controllers), managing AI/ML model deployment, "
        "distributing configuration files, and analyzing operational logs. This document includes a "
        "project timeline with weekly task scheduling from March 2026 to June 2026."
    )

    doc.add_heading("1.2 Scope", level=2)
    doc.add_paragraph(
        "The Factory Monitoring System is a 4-component distributed application deployed on a factory "
        "Local Area Network (LAN). It consists of:"
    )
    for item in [
        "A C++ Agent application running on each factory PC (Machine Controller).",
        "An ASP.NET Core Web API backend hosted on a central server.",
        "A React-based web dashboard accessed via browser by factory engineers.",
        "A SQL Server database for persistent data storage.",
    ]:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading("1.3 Intended Audience", level=2)
    tbl(doc, ["Audience", "Usage"],
        [["Project Lead", "Milestone tracking, feature sign-off, resource planning"],
         ["Developers", "Implementation reference, system specs, dependency management"],
         ["IT / DevOps", "Deployment, hosting requirements, infrastructure planning"],
         ["QA / Testers", "Functional and non-functional test cases"]],
        widths=[4, 13])

    doc.add_heading("1.4 Project Team", level=2)
    tbl(doc, ["Name", "Role", "Modules Owned"],
        [["Divyansh", "Intern — Model Management & UI",
          "Dashboard, model distribution, remote config, UI/UX, sidebar, theme"],
         ["Bhavik", "Intern — Log Analyser",
          "Log parsing, Gantt charts, barrel execution, inspection images"],
         ["Praganya", "Intern — Model Library & UI",
          "Model Library, upload/download, versioning, diff viewer, XML editor"]],
        widths=[3, 5, 9])

    doc.add_heading("1.5 Definitions & Abbreviations", level=2)
    tbl(doc, ["Term", "Definition"],
        [["MC", "Machine Controller — A factory PC on a production line"],
         ["Line", "A production line containing multiple MCs"],
         ["Model Version", "Software/hardware version of PCs (e.g., 3.5, 4.0)"],
         ["Model (AI)", "Machine learning model package (.zip) distributed to MCs"],
         ["Agent", "C++ application on each MC for server communication"],
         ["Heartbeat", "Periodic status signal from Agent to Server"],
         ["Target Model", "Designated standard model for a production line"],
         ["NG", "No Good — Factory quality term for failed inspection"],
         ["SRS", "Software Requirements Specification"],
         ["LAN", "Local Area Network"]],
        widths=[4, 13])

    doc.add_heading("1.6 References", level=2)
    for r in ["IEEE 830-1998 Standard for SRS", "Project README.md",
              "Database Scripts: db/01_CreateDatabase.sql – db/06_UpdateSampleData.sql",
              "Microsoft .NET 8 Documentation", "React 18 Documentation"]:
        doc.add_paragraph(r, style='List Bullet')

    doc.add_page_break()

    # ═══ 2. OVERALL DESCRIPTION ═══
    doc.add_heading("2. Overall Description", level=1)

    doc.add_heading("2.1 Product Perspective", level=2)
    doc.add_paragraph(
        "The Factory Monitoring System is a standalone, custom-built enterprise application deployed "
        "entirely within a factory's private LAN. It replaces manual model deployment (via USB drives), "
        "manual configuration editing (via remote desktop), and manual log analysis (via file copying) "
        "with a centralized, web-based, automated solution."
    )

    doc.add_heading("2.2 Product Functions", level=2)
    funcs = [
        ("Real-Time PC Monitoring", "Live dashboard showing online/offline status, application heartbeat, model compliance for all factory PCs grouped by production line."),
        ("Remote Configuration", "Upload, download, and push changes to PC configuration files (.ini, .json) remotely through the web interface."),
        ("Model Library", "Central repository for ML model packages (.zip) with metadata, search, categorization."),
        ("Model Distribution", "Deploy models to specific PCs, entire lines, specific versions, or all PCs — with compliance tracking."),
        ("Model Editor", "In-browser code editor for viewing/editing files inside model archives with syntax highlighting, diff view, undo/redo."),
        ("Version History", "Track all model edits with version snapshots, side-by-side diff comparison, and rollback capability."),
        ("Log Analytics", "Parse factory operation logs into interactive Gantt chart timelines, barrel execution breakdowns, and inspection image galleries."),
        ("Agent Communication", "Automated agent registration, heartbeat, model/config sync, log/image upload, and remote command execution."),
    ]
    for name, desc in funcs:
        p = doc.add_paragraph()
        r = p.add_run(f"{name}: "); r.font.bold = True
        p.add_run(desc)

    doc.add_heading("2.3 User Classes", level=2)
    tbl(doc, ["User Class", "Description", "Access Level"],
        [["Factory Operator", "Monitors dashboard, views PC status", "Read-only dashboard"],
         ["Factory Engineer", "Uploads models, edits configs, analyzes logs", "Full read/write"],
         ["Administrator", "Manages PCs, lines, system settings", "Full admin"]],
        widths=[4, 8, 5])

    doc.add_heading("2.4 Constraints & Assumptions", level=2)
    doc.add_paragraph("Constraints:", style='List Bullet')
    for c in [
        "All components must run on Windows (Agent uses Win32 API).",
        "Database storage only — no network file shares required.",
        "System operates on isolated factory LAN (no internet required after initial setup).",
        "No authentication/authorization currently implemented (trusted network assumption).",
    ]:
        p = doc.add_paragraph(c)
        p.paragraph_format.left_indent = Cm(1.5)

    doc.add_page_break()

    # ═══ 3. SYSTEM REQUIREMENTS (kept same) ═══
    doc.add_heading("3. System Requirements", level=1)
    doc.add_paragraph("This section specifies the hardware and software requirements for each component.")

    doc.add_heading("3.1 Server (Hosting) PC Requirements", level=2)
    doc.add_paragraph("The server PC hosts both the ASP.NET Core Web API and the SQL Server database.")
    doc.add_heading("Hardware", level=3)
    tbl(doc, ["Component", "Minimum", "Recommended"],
        [["Processor", "Any modern multi-core CPU (4+ cores)", "8+ core CPU"],
         ["RAM", "8 GB", "16 GB or more"],
         ["Storage", "50 GB free (SSD preferred)", "100 GB+ SSD"],
         ["Network", "1 Gbps Ethernet (LAN)", "1 Gbps Ethernet (LAN)"]],
        widths=[3, 6.5, 7.5])

    doc.add_heading("3.2 Factory Agent PC Requirements", level=2)
    tbl(doc, ["Component", "Minimum", "Notes"],
        [["Processor", "Any x86/x64 CPU", "Agent is lightweight (~5 MB binary)"],
         ["RAM", "2 GB", "Agent uses < 50 MB"],
         ["Storage", "500 MB free", "For agent binary, models, config, logs"],
         ["Network", "100 Mbps Ethernet", "Must reach server IP"]],
        widths=[3, 5, 9])

    doc.add_heading("3.3 Client (Dashboard) PC Requirements", level=2)
    tbl(doc, ["Component", "Requirement"],
        [["Browser", "Chrome 90+, Edge 90+, or Firefox 90+"],
         ["Network", "LAN connectivity to the server PC"],
         ["Display", "1280×720 min (1920×1080 recommended)"],
         ["RAM", "4 GB minimum"]],
        widths=[3, 14])

    doc.add_heading("3.4 Network Requirements", level=2)
    tbl(doc, ["Requirement", "Details"],
        [["Topology", "All components on the same factory LAN"],
         ["Bandwidth", "100 Mbps sufficient (model files up to 10 MB)"],
         ["Ports", "Server API: 5000 (HTTP). Dashboard dev: 5173"],
         ["Internet", "Not required after initial software installation"]],
        widths=[4, 13])

    doc.add_page_break()

    # ═══ 4. SYSTEM ARCHITECTURE (brief) ═══
    doc.add_heading("4. System Architecture", level=1)
    tbl(doc, ["Component", "Technology", "Deployment", "Port"],
        [["Factory Agent", "C++20 / Win32 / WinHTTP", "Each factory MC (system tray app)", "N/A"],
         ["Web API", "ASP.NET Core 8.0 / EF Core", "Central server PC", "5000"],
         ["Dashboard", "React 18 / TypeScript / Vite", "Browser", "5173 (dev)"],
         ["Database", "SQL Server 2019+", "Central server PC", "1433"]],
        widths=[3, 5, 5.5, 3.5])

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════
    # 5. FUNCTIONAL REQUIREMENTS — GANTT CHART TIMELINE
    # ═══════════════════════════════════════════════════════
    doc.add_heading("5. Functional Requirements — Project Timeline", level=1)
    doc.add_paragraph(
        "The following Gantt chart tables show the weekly schedule for implementing all functional "
        "requirements, organized by module. Timeline: March 2026 (W1) to June 2026 (W4)."
    )

    # Module 1: Real-Time Monitoring Dashboard
    # (functionality, main_task, sub_task, start_week_idx, end_week_idx)
    # Weeks: 0-3=Mar, 4-7=Apr, 8-11=May, 12-15=Jun
    create_gantt_table(doc, "5.1 Real-Time Monitoring Dashboard", [
        ("Monitoring", "Dashboard Layout", "PC grouping by production line", 0, 1),
        ("Monitoring", "Status Indicators", "Online/offline color-coded indicators", 0, 1),
        ("Monitoring", "Application Status", "App running status display", 1, 2),
        ("Monitoring", "Version Filtering", "Filter PCs by model version via sidebar", 2, 3),
        ("Monitoring", "View Modes", "Grid and List view support", 2, 3),
        ("Monitoring", "Line Accordion", "Collapsible/expandable production lines", 3, 4),
        ("Monitoring", "Compliance Badge", "Model compliance status per line", 3, 4),
        ("Monitoring", "Statistics", "Aggregate stats (total, online, offline)", 4, 5),
        ("Monitoring", "PC Card Design", "IP display, model name, card modal", 4, 5),
        ("Monitoring", "Manage Models", "Per-line model management button", 5, 6),
        ("Monitoring", "Testing & QA", "Integration testing & bug fixes", 6, 7),
    ], color_start=0)

    # Module 2: PC Details & Management
    create_gantt_table(doc, "5.2 PC Details & Management", [
        ("PC Management", "PC Detail View", "Full PC info modal (line, IP, version, paths)", 1, 2),
        ("PC Management", "Model Operations", "Upload model to PC", 2, 3),
        ("PC Management", "Model Operations", "Download model from PC", 3, 4),
        ("PC Management", "Model Operations", "Change active model on PC", 3, 4),
        ("PC Management", "Model Operations", "Delete model from PC", 4, 5),
        ("PC Management", "Config Operations", "Download config file from PC", 4, 5),
        ("PC Management", "Config Operations", "Upload new config to PC", 5, 6),
        ("PC Management", "PC Admin", "Edit PC metadata", 5, 6),
        ("PC Management", "PC Admin", "Delete PC registration", 6, 6),
        ("PC Management", "Download Progress", "Poll & display download status", 6, 7),
        ("PC Management", "Testing & QA", "End-to-end testing", 7, 8),
    ], color_start=2)

    # Module 3: Model Library
    create_gantt_table(doc, "5.3 Model Library", [
        ("Model Library", "Repository Page", "Centralized model library page UI", 2, 3),
        ("Model Library", "Upload System", "Upload ZIP models with metadata", 3, 4),
        ("Model Library", "Display & Table", "Model table with name, size, date, category", 3, 5),
        ("Model Library", "Download", "Download model files from library", 4, 5),
        ("Model Library", "Delete", "Delete models from library", 5, 5),
        ("Model Library", "Search", "Search/filter models by name", 5, 6),
        ("Model Library", "Testing & QA", "Integration testing", 6, 7),
    ], color_start=4)

    # Module 4: Model Distribution & Deployment
    create_gantt_table(doc, "5.4 Model Distribution & Deployment", [
        ("Distribution", "Distribution Engine", "Distribute to all / version / line / PCs", 4, 6),
        ("Distribution", "Target Model", "Set line target model on distribution", 5, 6),
        ("Distribution", "Compliance", "Model compliance tracking per line", 6, 7),
        ("Distribution", "Offline Handling", "Alert for offline PCs, proceed option", 7, 8),
        ("Distribution", "Line Delete", "Delete model from all PCs on line", 7, 8),
        ("Distribution", "PC Selection", "PC checklist for targeted deployment", 8, 9),
        ("Distribution", "Version Scope", "Version-scoped model operations", 8, 9),
        ("Distribution", "Testing & QA", "Distribution pipeline testing", 9, 10),
    ], color_start=1)

    doc.add_page_break()

    # Module 5: Model Editor
    create_gantt_table(doc, "5.5 Model Editor", [
        ("Model Editor", "Editor Core", "Open model in browser editor", 4, 5),
        ("Model Editor", "File Tree", "Display ZIP file tree structure", 5, 6),
        ("Model Editor", "Tab System", "Multi-file tabs with dirty indicators", 5, 6),
        ("Model Editor", "Syntax Highlight", "Highlighting for XML, JSON, Python, etc.", 6, 7),
        ("Model Editor", "Undo/Redo", "Undo and redo operations", 6, 7),
        ("Model Editor", "Unsaved Warning", "Prompt before navigating away", 7, 7),
        ("Model Editor", "Diff View", "Side-by-side diff with word-level highlight", 7, 8),
        ("Model Editor", "Save System", "Save changes back to ZIP in DB", 8, 9),
        ("Model Editor", "Bulk Save", "Save all modified files at once", 8, 9),
        ("Model Editor", "Diff Revert", "Revert individual lines in diff", 9, 10),
        ("Model Editor", "Testing & QA", "Editor workflow testing", 10, 11),
    ], color_start=3)

    # Module 6: XML Visual Editor
    create_gantt_table(doc, "5.6 XML Visual Editor", [
        ("XML Editor", "Parser", "Parse XML into Groups → Specs → Values tree", 7, 8),
        ("XML Editor", "Visual Tree", "Display structured parameter tree", 8, 9),
        ("XML Editor", "Inline Edit", "Inline editing of parameter values", 8, 9),
        ("XML Editor", "Inline Diff", "Show diffs for changed values", 9, 10),
        ("XML Editor", "Revert", "Revert individual parameter changes", 9, 10),
        ("XML Editor", "Expand Values", "Expand/collapse long values on click", 10, 10),
        ("XML Editor", "Testing & QA", "XML editor testing", 10, 11),
    ], color_start=5)

    # Module 7: Model Version History
    create_gantt_table(doc, "5.7 Model Version History", [
        ("Version History", "Auto Snapshot", "Auto-create version on save", 8, 9),
        ("Version History", "Timeline UI", "Version list with date, author, summary", 9, 10),
        ("Version History", "File Diff", "View file-level changes between versions", 9, 10),
        ("Version History", "Diff Viewer", "Side-by-side diff with minimap indicators", 10, 11),
        ("Version History", "Param Changes", "XML param changes table view", 10, 11),
        ("Version History", "Revert", "Revert model to any previous version", 11, 12),
        ("Version History", "Testing & QA", "Version history testing", 12, 13),
    ], color_start=7)

    doc.add_page_break()

    # Module 8: Log Analyser
    create_gantt_table(doc, "5.8 Log Analyser", [
        ("Log Analyser", "PC Selection", "Select PC and browse log folder", 2, 3),
        ("Log Analyser", "Log Fetching", "Fetch log files with caching", 3, 4),
        ("Log Analyser", "Log Parsing", "Parse operation data from logs", 3, 5),
        ("Log Analyser", "Gantt Chart", "Interactive Plotly.js Gantt timeline", 5, 7),
        ("Log Analyser", "Barrel Chart", "Barrel execution breakdown chart", 6, 8),
        ("Log Analyser", "Raw View", "View raw log file content", 7, 8),
        ("Log Analyser", "Images", "Inspection images for NG operations", 8, 9),
        ("Log Analyser", "Lazy Loading", "Lazy-load images from agents", 9, 10),
        ("Log Analyser", "Download", "Download log files", 9, 10),
        ("Log Analyser", "Thumbnails", "Thumbnail previews with hover tooltip", 10, 11),
        ("Log Analyser", "Deduplication", "Deduplicate concurrent log requests", 10, 11),
        ("Log Analyser", "Testing & QA", "Log analyser end-to-end testing", 11, 12),
    ], color_start=0)

    # Module 9: Agent Communication
    create_gantt_table(doc, "5.9 Agent Communication", [
        ("Agent", "Registration", "Agent registration with server", 0, 1),
        ("Agent", "Registration", "First-run registration dialog", 0, 1),
        ("Agent", "Heartbeat", "Periodic heartbeat with app status", 1, 2),
        ("Agent", "Commands", "Poll and execute pending commands", 2, 3),
        ("Agent", "Model Sync", "Sync model inventory with server", 2, 3),
        ("Agent", "Distribution", "Download pending model distributions", 3, 4),
        ("Agent", "Config Sync", "Upload config & apply pending updates", 4, 5),
        ("Agent", "Log Sync", "Sync log folder structure", 5, 6),
        ("Agent", "Log Upload", "Upload log content on demand", 5, 6),
        ("Agent", "Image Upload", "Upload inspection images (GZIP/multipart)", 6, 7),
        ("Agent", "Model Upload", "Upload model files to server", 7, 8),
        ("Agent", "System Tray", "Windows system tray application", 0, 0),
        ("Agent", "Testing & QA", "Agent communication testing", 8, 9),
    ], color_start=2)

    # Module 10: UI Framework & Navigation
    create_gantt_table(doc, "5.10 UI Framework & Navigation", [
        ("UI Framework", "Sidebar", "Collapsible sidebar with version tree", 0, 1),
        ("UI Framework", "Sidebar Stats", "Per-line online/offline counts", 1, 2),
        ("UI Framework", "Theme System", "Light and Dark theme toggle", 1, 2),
        ("UI Framework", "Routing", "Client-side routing for all pages", 2, 3),
        ("UI Framework", "404 Page", "Not Found page for invalid routes", 2, 3),
        ("UI Framework", "Quick Links", "Dashboard, Library, Log Analyzer links", 3, 3),
        ("UI Framework", "Polish & QA", "UI polish and cross-browser testing", 13, 15),
    ], color_start=6)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════
    # 6. NON-FUNCTIONAL REQUIREMENTS — GANTT CHART
    # ═══════════════════════════════════════════════════════
    doc.add_heading("6. Non-Functional Requirements — Project Timeline", level=1)
    doc.add_paragraph(
        "Non-functional requirements are implemented continuously throughout the project. "
        "The following timeline shows the primary implementation and validation periods."
    )

    create_gantt_table(doc, "6.1 Non-Functional Requirements Schedule", [
        ("Performance", "NFR-001", "Dashboard loads within 3s on LAN", 4, 7),
        ("Performance", "NFR-002", "Model upload supports up to 10 MB", 3, 5),
        ("Performance", "NFR-003", "LRU cache for log files (100 MB limit)", 6, 8),
        ("Performance", "NFR-004", "Image caching with configurable TTL", 8, 10),
        ("Usability", "NFR-005", "Consistent color-coding (green/red/amber)", 0, 3),
        ("Usability", "NFR-006", "Unique IDs for all interactive elements", 0, 15),
        ("Usability", "NFR-007", "Responsive UI for modern browsers", 0, 15),
        ("Usability", "NFR-008", "Light and Dark theme support", 1, 2),
        ("Reliability", "NFR-009", "Auto-mark PCs offline after missed heartbeat", 2, 4),
        ("Reliability", "NFR-010", "Structured API error responses", 0, 15),
        ("Reliability", "NFR-011", "10s connection timeout with graceful errors", 3, 5),
        ("Maintainability", "NFR-012", "Backend: Controllers/Services/Models separation", 0, 15),
        ("Scalability", "NFR-013", "Database indexes on key columns", 0, 2),
        ("Security", "NFR-014", "File path validation (anti-traversal)", 3, 5),
        ("Security", "NFR-015", "Trusted network model (factory LAN only)", 0, 0),
    ], color_start=0)

    doc.add_page_break()

    # ═══ 7. DATABASE SCHEMA SUMMARY ═══
    doc.add_heading("7. Database Schema Summary", level=1)
    doc.add_paragraph("Database: FactoryMonitoringDB — SQL Server. 9 tables total.")
    tbl(doc, ["Table", "Purpose", "Key Columns"],
        [["FactoryMCs", "PC registrations", "MCId, LineNumber, IPAddress, IsOnline"],
         ["ConfigFiles", "PC configuration content", "ConfigId, MCId, PendingUpdate"],
         ["Models", "Models discovered on PCs", "ModelId, MCId, IsCurrentModel"],
         ["ModelFiles", "Centralized model repository", "ModelFileId, FileData (ZIP)"],
         ["ModelVersions", "Version history snapshots", "ModelVersionId, VersionNumber"],
         ["LineTargetModels", "Target model per line+version", "LineNumber, ModelVersion"],
         ["ModelDistributions", "Deployment tracking", "DistributionId, Status"],
         ["AgentCommands", "Remote command queue", "CommandId, CommandType, Status"],
         ["SystemLogs", "Audit trail", "LogId, Action, Timestamp"]],
        widths=[3.5, 4, 9.5])

    doc.add_page_break()

    # ═══ 8. API ENDPOINT SUMMARY ═══
    doc.add_heading("8. API Endpoint Summary", level=1)
    doc.add_paragraph("Base URL: http://<server-ip>:5000/api — Total: 50+ endpoints.")

    api_groups = [
        ("Dashboard & PC Management (7 endpoints)", [
            ["GET", "/api/pcs", "Get PCs (filter by version/line)"],
            ["GET", "/api/stats", "Aggregate statistics"],
            ["PUT", "/api/mc/{id}", "Update PC metadata"],
            ["DELETE", "/api/mc/{id}", "Delete PC"],
        ]),
        ("Model Library (10 endpoints)", [
            ["POST", "/api/modellibrary/upload", "Upload model"],
            ["GET", "/api/modellibrary/{id}/download", "Download model"],
            ["POST", "/api/modellibrary/apply", "Distribute model"],
            ["POST", "/api/modellibrary/{id}/files/save", "Bulk save files"],
        ]),
        ("Agent Communication (12 endpoints)", [
            ["POST", "/api/agent/register", "Register agent"],
            ["POST", "/api/agent/heartbeat", "Heartbeat + commands"],
            ["POST", "/api/agent/syncmodels", "Sync model list"],
            ["POST", "/api/agent/uploadimagesbinary", "Upload images"],
        ]),
        ("Log Analyser (7 endpoints)", [
            ["POST", "/api/loganalyzer/analyze", "Parse log file"],
            ["GET", "/api/loganalyzer/images/{mcId}", "Inspection images"],
        ]),
    ]

    for title, rows in api_groups:
        doc.add_heading(title, level=2)
        tbl(doc, ["Method", "Endpoint", "Description"], rows, widths=[2, 7.5, 7.5])
        doc.add_paragraph()

    doc.add_page_break()

    # ═══ 9. DEPLOYMENT ═══
    doc.add_heading("9. Deployment & Installation Guide", level=1)

    doc.add_heading("9.1 Database Setup", level=2)
    for s in [
        "1. Install SQL Server 2019/2022 (Express edition is sufficient).",
        "2. Execute database scripts in order: 01_CreateDatabase → 02_CreateTables.",
        "3. Verify 9 tables are created.",
    ]:
        doc.add_paragraph(s)

    doc.add_heading("9.2 Backend Server Setup", level=2)
    for s in [
        "1. Install .NET 8.0 Runtime.",
        "2. Edit appsettings.json with SQL Server connection string.",
        '3. Run: dotnet run --urls "http://0.0.0.0:5000"',
        "4. Open firewall port 5000.",
    ]:
        doc.add_paragraph(s)

    doc.add_heading("9.3 Agent Deployment", level=2)
    for s in [
        "1. Build FactoryAgent.sln (Release x64).",
        "2. Copy FactoryAgent.exe to each factory PC.",
        "3. Run — registration dialog appears on first launch.",
    ]:
        doc.add_paragraph(s)

    doc.add_page_break()

    # ═══ 10. GLOSSARY ═══
    doc.add_heading("10. Glossary", level=1)
    tbl(doc, ["Term", "Definition"],
        [["MC", "Machine Controller — A factory PC on a production line"],
         ["Line", "Production line containing multiple MCs"],
         ["Model Version", "Software/hardware version string (e.g., '3.5', '4.0')"],
         ["Model", "Machine learning model package stored as a ZIP archive"],
         ["Target Model", "Designated standard model for a production line"],
         ["Compliance", "State where PC's active model matches line's target"],
         ["Agent", "C++ desktop application running on each factory PC"],
         ["Heartbeat", "Periodic HTTP call from agent to server"],
         ["NG", "No Good — Failed inspection result"],
         ["LAN", "Local Area Network"],
         ["SRS", "Software Requirements Specification"]],
        widths=[4, 13])

    # Footer
    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("— End of Document —"); r.font.italic = True
    r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Factory Monitoring System — SRS v1.0 — February 2026")
    r.font.size = Pt(10); r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.save(OUT)
    print(f"Done → {OUT}")

if __name__ == "__main__":
    build()
